import streamlit as st
import pandas as pd
import numpy as np
import yfinance as yf
from bcb import sgs
import re
import io
import requests
import plotly.express as px
import json
import os
import sqlite3
import hashlib

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ==========================================
# 1. CONFIGURAÇÃO E FORMATADORES
# ==========================================
st.set_page_config(page_title="Terminal de Gestão CNPI", layout="wide")

def f_brl(x): return f"R$ {float(x):,.2f}".replace(",", "v").replace(".", ",").replace("v", ".")
def f_brl_4(x): return f"R$ {float(x):,.4f}".replace(",", "v").replace(".", ",").replace("v", ".")
def f_pct(x): return f"{float(x):,.2f}%".replace(",", "v").replace(".", ",").replace("v", ".")

# Injetado aqui no topo para eliminar o NameError definitivo
def limpar_numero(x):
    if pd.isna(x): return 0.0
    if isinstance(x, (int, float, np.number)): return float(x)
    try: return float(str(x).replace('R$', '').replace('.', '').replace(',', '.').strip())
    except: return 0.0

def to_excel(df, sheet_name='Sheet1'):
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name=sheet_name)
        worksheet = writer.sheets[sheet_name]
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        header_font = Font(name="Arial", size=11, bold=True, color="FFFFFF")
        thin_border = Border(
            left=Side(style='thin', color='D9D9D9'), 
            right=Side(style='thin', color='D9D9D9'), 
            top=Side(style='thin', color='D9D9D9'), 
            bottom=Side(style='thin', color='D9D9D9')
        )
        
        for col_num, column in enumerate(worksheet.columns, 1):
            cell = worksheet.cell(row=1, column=col_num)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
            
            for row_num in range(2, worksheet.max_row + 1):
                data_cell = worksheet.cell(row=row_num, column=col_num)
                data_cell.font = Font(name="Arial", size=10)
                data_cell.border = thin_border
                
                if isinstance(data_cell.value, (int, float)): 
                    data_cell.alignment = Alignment(horizontal="right")
                else: 
                    data_cell.alignment = Alignment(horizontal="center")
                    
            max_len = max(len(str(c.value or '')) for c in column)
            col_letter = get_column_letter(col_num)
            worksheet.column_dimensions[col_letter].width = max(max_len + 4, 13)
            
    return output.getvalue()

def export_docx(historico):
    if not HAS_DOCX: 
        return None
        
    doc = docx.Document()
    doc.add_heading('Relatório Analítico Institucional - Comitê de IA CNPI', 0)
    
    for msg in historico:
        if msg["role"] == "user":
            doc.add_heading("Consulta Operacional:", level=2)
        else:
            doc.add_heading("Parecer da Gestora IA:", level=2)
            
        doc.add_paragraph(msg["content"])
        
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

MAPEAMENTO_TICKERS = {
    "GALG11": "GARE11", "SOMA3": "ALOS3", "ARZZ3": "ALOS3", 
    "VVAR3": "BHIA3", "VIIA3": "BHIA3", "BRML3": "ALSO3", 
    "BBRK11": "BRCR11", "HCTR11": "TRXD11", "TORD11": "TRXD11"
}

UNITS_ACOES = [
    'SANB11', 'TAEE11', 'KLBN11', 'BPAC11', 'ALUP11', 
    'ENGI11', 'BIDI11', 'CPLE11', 'SAPR11', 'RNEW11'
]

# Inicialização de Variáveis de Sessão para garantir a persistência na navegação
if 'df_base' not in st.session_state: 
    st.session_state.df_base = pd.DataFrame()

if 'df_transacoes' not in st.session_state:
    st.session_state.df_transacoes = pd.DataFrame(columns=["Ativo", "Tipo", "Quantidade", "Preço Unitário", "Preço Médio na Época", "Resultado Realizado", "Data"])
    
if 'df_tesouro' not in st.session_state: 
    st.session_state.df_tesouro = pd.DataFrame()
    
if 'dados_mercado' not in st.session_state: 
    st.session_state.dados_mercado = {}
    
if 'df_simul' not in st.session_state: 
    st.session_state.df_simul = pd.DataFrame()
    
if 'logged_in' not in st.session_state: 
    st.session_state.logged_in = False
    
if 'username' not in st.session_state: 
    st.session_state.username = ""
    
if 'historico_chat' not in st.session_state: 
    st.session_state.historico_chat = []

# ==========================================
# 2. MOTOR DE BANCO DE DADOS HÍBRIDO (SQLite / PostgreSQL)
# ==========================================
IS_POSTGRES = "POSTGRES_URL" in st.secrets
PARAM = "%s" if IS_POSTGRES else "?"

def get_db_connection():
    if IS_POSTGRES:
        import psycopg2
        return psycopg2.connect(st.secrets["POSTGRES_URL"])
    else:
        return sqlite3.connect("terminal_cnpi.db")

def init_db():
    conn = get_db_connection()
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS usuarios (username TEXT PRIMARY KEY, password TEXT)''')
    
    if IS_POSTGRES:
        c.execute('''CREATE TABLE IF NOT EXISTS carteiras (username TEXT, ativo TEXT, quantidade DOUBLE PRECISION, preco_medio DOUBLE PRECISION, data_media TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS tesouro_v2 (username TEXT, titulo TEXT, data_compra TEXT, tipo_taxa TEXT, investido DOUBLE PRECISION, taxa DOUBLE PRECISION, vencimento INTEGER)''')
        c.execute('''CREATE TABLE IF NOT EXISTS proventos_ledger (username TEXT, ativo TEXT, data_ex TEXT, valor DOUBLE PRECISION)''')
        c.execute('''CREATE TABLE IF NOT EXISTS transacoes_ledger (username TEXT, ativo TEXT, tipo TEXT, quantidade DOUBLE PRECISION, preco DOUBLE PRECISION, preco_medio DOUBLE PRECISION, resultado DOUBLE PRECISION, data TEXT)''')
    else:
        c.execute('''CREATE TABLE IF NOT EXISTS carteiras (username TEXT, Ativo TEXT, Quantidade REAL, Preco_Medio REAL, Data_Media TEXT)''')
        c.execute('''CREATE TABLE IF NOT EXISTS tesouro_v2 (username TEXT, titulo TEXT, data_compra TEXT, tipo_taxa TEXT, investido REAL, taxa REAL, vencimento INTEGER)''')
        c.execute('''CREATE TABLE IF NOT EXISTS proventos_ledger (username TEXT, ativo TEXT, data_ex TEXT, valor REAL)''')
        c.execute('''CREATE TABLE IF NOT EXISTS transacoes_ledger (username TEXT, ativo TEXT, tipo TEXT, quantidade REAL, preco REAL, preco_medio REAL, resultado REAL, data TEXT)''')
        
    c.execute('''CREATE TABLE IF NOT EXISTS chat_ia (username TEXT, role TEXT, content TEXT)''')
    conn.commit()
    conn.close()

def hash_password(password): 
    return hashlib.sha256(password.encode()).hexdigest()

def registrar_usuario(username, password):
    # ==========================================
# BYPASS DE EMERGÊNCIA - RECRIAR USUÁRIO
# ==========================================
registrar_usuario("rodrigo", "admin123")
    conn = get_db_connection()
    c = conn.cursor()
    try:
        c.execute(f"INSERT INTO usuarios (username, password) VALUES ({PARAM}, {PARAM})", (username.strip(), hash_password(password)))
        conn.commit()
        conn.close()
        return True
    except:
        conn.close()
        return False

def autenticar_usuario(username, password):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute(f"SELECT * FROM usuarios WHERE username={PARAM} AND password={PARAM}", (username.strip(), hash_password(password)))
    user = c.fetchone()
    conn.close()
    return user is not None

def atualizar_senha(username, nova_senha):
    usr = username.strip()
    conn = get_db_connection()
    c = conn.cursor()
    c.execute(f"SELECT * FROM usuarios WHERE username={PARAM}", (usr,))
    user = c.fetchone()
    if user is None:
        conn.close()
        return False
        
    c.execute(f"UPDATE usuarios SET password={PARAM} WHERE username={PARAM}", (hash_password(nova_senha), usr))
    conn.commit()
    conn.close()
    return True

def salvar_dados_completos_db(username):
    conn = get_db_connection()
    c = conn.cursor()
    usr = username.strip()
    
    # Salvar Carteira B3
    c.execute(f"DELETE FROM carteiras WHERE username={PARAM}", (usr,))
    if not st.session_state.df_base.empty:
        for _, r in st.session_state.df_base.iterrows():
            c.execute(f"INSERT INTO carteiras (username, ativo, quantidade, preco_medio, data_media) VALUES ({PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM})",
                      (usr, str(r.get('Ativo', '')), float(limpar_numero(r.get('Quantidade', 0))), float(limpar_numero(r.get('Preço Médio', 0))), str(r.get('Data Média', ''))))
            
    # Salvar Tesouro
    c.execute(f"DELETE FROM tesouro_v2 WHERE username={PARAM}", (usr,))
    if isinstance(st.session_state.df_tesouro, pd.DataFrame) and not st.session_state.df_tesouro.empty:
        for _, r in st.session_state.df_tesouro.iterrows():
            if pd.isna(r.get('Título')) or str(r.get('Título')).strip() == '': 
                continue
            c.execute(f"INSERT INTO tesouro_v2 (username, titulo, data_compra, tipo_taxa, investido, taxa, vencimento) VALUES ({PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM})",
                      (usr, str(r.get('Título', 'Tesouro')), str(r.get('Data Compra', '')), str(r.get('Tipo Taxa', 'Pré-fixado')), float(limpar_numero(r.get('Valor Investido (R$)', 0))), float(limpar_numero(r.get('Taxa Contratada (%)', 0))), int(limpar_numero(r.get('Ano Vencimento', 2030)))))
            
    # Salvar Livro Razão de Proventos
    c.execute(f"DELETE FROM proventos_ledger WHERE username={PARAM}", (usr,))
    if 'df_ledger' in st.session_state and not st.session_state.df_ledger.empty:
        for _, r in st.session_state.df_ledger.iterrows():
            c.execute(f"INSERT INTO proventos_ledger (username, ativo, data_ex, valor) VALUES ({PARAM}, {PARAM}, {PARAM}, {PARAM})",
                      (usr, str(r['Ativo']), str(r['Data Ex']), float(r['Valor Recebido (R$)'])))

    # --- SALVAR LEDGER DE TRANSAÇÕES COMPRA/VENDA ---
    c.execute(f"DELETE FROM transacoes_ledger WHERE username={PARAM}", (usr,))
    if 'df_transacoes' in st.session_state and not st.session_state.df_transacoes.empty:
        for _, r in st.session_state.df_transacoes.iterrows():
            c.execute(f"INSERT INTO transacoes_ledger (username, ativo, tipo, quantidade, preco, preco_medio, resultado, data) VALUES ({PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM}, {PARAM})",
                      (usr, str(r['Ativo']), str(r['Tipo']), float(r['Quantidade']), float(r['Preço Unitário']), float(r['Preço Médio na Época']), float(r['Resultado Realizado']), str(r['Data'])))

    # Salvar Chat IA
    c.execute(f"DELETE FROM chat_ia WHERE username={PARAM}", (usr,))
    for msg in st.session_state.historico_chat[-30:]:
        c.execute(f"INSERT INTO chat_ia (username, role, content) VALUES ({PARAM}, {PARAM}, {PARAM})", (usr, msg['role'], msg['content']))
        
    conn.commit()
    conn.close()

def carregar_dados_completos_db(username):
    conn = get_db_connection()
    c = conn.cursor()
    usr = username.strip()
    
    c.execute(f"SELECT ativo, quantidade, preco_medio, data_media FROM carteiras WHERE username={PARAM}", (usr,))
    df_cart = pd.DataFrame(c.fetchall(), columns=["Ativo", "Quantidade", "Preço Médio", "Data Média"])
    if not df_cart.empty: 
        df_cart['Data Média'] = pd.to_datetime(df_cart['Data Média']).dt.date
    else:
        df_cart = pd.DataFrame(columns=["Ativo", "Quantidade", "Preço Médio", "Data Média"])
    st.session_state.df_base = df_cart
    
    c.execute(f"SELECT titulo, data_compra, tipo_taxa, investido, taxa, vencimento FROM tesouro_v2 WHERE username={PARAM}", (usr,))
    df_tes = pd.DataFrame(c.fetchall(), columns=["Título", "Data Compra", "Tipo Taxa", "Valor Investido (R$)", "Taxa Contratada (%)", "Ano Vencimento"])
    if not df_tes.empty: 
        df_tes['Data Compra'] = pd.to_datetime(df_tes['Data Compra']).dt.date
    else:
        df_tes = pd.DataFrame(columns=["Título", "Data Compra", "Tipo Taxa", "Valor Investido (R$)", "Taxa Contratada (%)", "Ano Vencimento", "Valor Futuro no Vencimento"])
    st.session_state.df_tesouro = df_tes

    c.execute(f"SELECT ativo, data_ex, valor FROM proventos_ledger WHERE username={PARAM}", (usr,))
    df_led = pd.DataFrame(c.fetchall(), columns=["Ativo", "Data Ex", "Valor Recebido (R$)"])
    if not df_led.empty: 
        df_led['Data Ex'] = pd.to_datetime(df_led['Data Ex']).dt.date
    st.session_state.df_ledger = df_led

    # --- CARREGAR LEDGER DE TRANSAÇÕES COMPRA/VENDA ---
    c.execute(f"SELECT ativo, tipo, quantidade, preco, preco_medio, resultado, data FROM transacoes_ledger WHERE username={PARAM}", (usr,))
    df_tx = pd.DataFrame(c.fetchall(), columns=["Ativo", "Tipo", "Quantidade", "Preço Unitário", "Preço Médio na Época", "Resultado Realizado", "Data"])
    if not df_tx.empty: 
        df_tx['Data'] = pd.to_datetime(df_tx['Data']).dt.date
    st.session_state.df_transacoes = df_tx
    
    c.execute(f"SELECT role, content FROM chat_ia WHERE username={PARAM}", (usr,))
    df_chat = pd.DataFrame(c.fetchall(), columns=["role", "content"])
    conn.close()
    
    if not df_chat.empty: 
        st.session_state.historico_chat = df_chat.to_dict('records')
    else: 
        st.session_state.historico_chat = [{"role": "assistant", "content": f"Saudações, {username}. O terminal está mapeado e pronto."}]

init_db()

# ==========================================
# 3. GATEKEEPER - TELA DE AUTENTICAÇÃO
# ==========================================
if not st.session_state.logged_in:
    st.markdown("<h1 style='text-align: center;'>🔐 Terminal de Gestão Profissional</h1>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Acesso restrito. Identifique-se para carregar seu portfólio.</p>", unsafe_allow_html=True)
    
    col_log1, col_log2, col_log3 = st.columns([1, 1, 1])
    
    with col_log2:
        tab_login, tab_register, tab_forgot = st.tabs(["Acesso", "Novo Registro", "Recuperar Senha"])
        
        with tab_login:
            login_user = st.text_input("Usuário", key="log_user").strip()
            login_pass = st.text_input("Senha", type="password", key="log_pass")
            
            if st.button("Entrar no Terminal", use_container_width=True):
                if autenticar_usuario(login_user, login_pass):
                    st.session_state.logged_in = True
                    st.session_state.username = login_user
                    carregar_dados_completos_db(login_user)
                    st.rerun()
                else: 
                    st.error("Credenciais inválidas.")
                
        with tab_register:
            reg_user = st.text_input("Novo Usuário", key="reg_user").strip()
            reg_pass = st.text_input("Nova Senha", type="password", key="reg_pass")
            
            if st.button("Registrar Acesso", use_container_width=True):
                if reg_user and reg_pass:
                    if registrar_usuario(reg_user, reg_pass): 
                        st.success("Conta provisionada com sucesso! Você já pode fazer o login.")
                    else: 
                        st.error("O nome de usuário já está em uso no sistema.")
                else: 
                    st.warning("Preencha ambos os campos corretamente.")
                
        with tab_forgot:
            forgot_user = st.text_input("Usuário Cadastrado", key="for_user").strip()
            forgot_pass = st.text_input("Nova Senha Desejada", type="password", key="for_pass")
            
            if st.button("Redefinir Senha", use_container_width=True):
                if atualizar_senha(forgot_user, forgot_pass): 
                    st.success("Sua senha foi redefinida com sucesso.")
                else: 
                    st.error("Usuário não encontrado no banco de dados.")
                    
    st.stop()
    # ==========================================
# 4. FUNÇÕES DE PROCESSAMENTO B3 E MACRO
# ==========================================
st.markdown(f"""
    <div style="text-align: center; margin-bottom: 20px;">
        <h3 style="font-weight: 400; margin-bottom: 0;">💼 Terminal de Gestão</h3>
        <h6 style="color: #666; font-weight: 300;">Analista Sênior: {st.session_state.username.upper()}</h6>
    </div>
""", unsafe_allow_html=True)
st.write("---")

@st.cache_data(ttl=86400)
def obter_macro_atual():
    selic_atual = 10.50
    ipca_12m = 4.00
    
    try:
        res = requests.get("https://brasilapi.com.br/api/taxas/v1", timeout=5)
        if res.status_code == 200:
            for taxa in res.json():
                if taxa['nome'] == 'Selic': 
                    selic_atual = float(taxa['valor'])
    except:
        try:
            selic_df = sgs.get({'selic': 432}, last=1)
            if not selic_df.empty: 
                selic_atual = float(selic_df['selic'].iloc[-1])
        except: 
            pass
            
    try:
        ipca_df = sgs.get({'IPCA_12M': 13522}, last=1)
        if not ipca_df.empty: 
            ipca_12m = float(ipca_df['IPCA_12M'].iloc[-1])
    except: 
        pass
        
    return selic_atual, ipca_12m

@st.cache_data(ttl=86400)
def carregar_dados_mercado():
    macro = pd.DataFrame()
    fundamentos = {}
    
    try:
        macro = sgs.get({'CDI': 12, 'IPCA': 433}, start='2019-01-01')
        macro['CDI'] = macro['CDI'] / 100
        macro['IPCA'] = macro['IPCA'] / 100
        
        url_fundamentus = 'https://www.fundamentus.com.br/resultado.php'
        headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
        df_fund = pd.read_html(io.StringIO(requests.get(url_fundamentus, headers=headers, timeout=10).text), decimal=',', thousands='.')[0]
        
        for _, row in df_fund.iterrows():
            t = str(row['Papel']).strip().upper()
            c = float(row['Cotação'])
            pl = float(row['P/L'])
            pvp = float(row['P/VP'])
            fundamentos[t] = {
                'vpa': c/pvp if pvp > 0 else 0.0, 
                'lpa': c/pl if pl > 0 else 0.0
            }
    except Exception as e: 
        pass
    
    ano_at = pd.Timestamp.now().year
    proj = {}
    
    try:
        url_focus = "https://olinda.bcb.gov.br/olinda/servico/Expectativas/versao/v1/odata/ExpectativasMercadoAnuais?$top=300&$filter=Indicador%20eq%20'IPCA'%20or%20Indicador%20eq%20'Selic'&$orderby=Data%20desc&$format=json"
        df_p = pd.DataFrame(requests.get(url_focus, timeout=5).json().get('value', []))
        
        if not df_p.empty:
            df_p = df_p[df_p['Data'] == df_p['Data'].max()]
            for ao in [0, 1, 2]:
                aa = str(ano_at + ao)
                df_a = df_p[df_p['DataReferencia'] == aa]
                if not df_a[df_a['Indicador'] == 'IPCA'].empty: 
                    proj[f"IPCA_{aa}"] = float(df_a[df_a['Indicador'] == 'IPCA']['Mediana'].values[0])
                if not df_a[df_a['Indicador'] == 'Selic'].empty: 
                    proj[f"Selic_{aa}"] = float(df_a[df_a['Indicador'] == 'Selic']['Mediana'].values[0])
    except: 
        pass
        
    return macro, fundamentos, proj, ano_at

selic_hoje, ipca_12m_hoje = obter_macro_atual()
df_macro, fundamentos_br, proj_focus, ano_atual = carregar_dados_mercado()

def limpar_numero(x):
    if pd.isna(x): 
        return 0.0
    if isinstance(x, (int, float, np.number)): 
        return float(x)
    try: 
        return float(str(x).replace('R$', '').replace('.', '').replace(',', '.').strip())
    except: 
        return 0.0

def traduzir_setor(setor_en):
    dicionario = {
        "Banks": "Bancos", "Utilities - Regulated Electric": "Energia", 
        "Real Estate - Retail": "Shoppings/Varejo", "REIT - Retail": "Shoppings/Varejo", 
        "Real Estate - Industrial": "Logística", "REIT - Industrial": "Logística", 
        "REIT - Office": "Lajes Corporativas", "REIT - Diversified": "Fundo Híbrido", 
        "Financial Data & Stock Exchanges": "Bolsa de Valores", 
        "Insurance": "Seguradoras", "Oil & Gas Integrated": "Petróleo e Gás"
    }
    return dicionario.get(setor_en, "Outros Setores")

def consolidar_carteira(df):
    if df.empty: 
        return df
        
    df['Ativo'] = df['Ativo'].astype(str).str.strip().str.upper().apply(lambda x: MAPEAMENTO_TICKERS.get(x, x))
    linhas = []
    
    for ativo, group in df.groupby('Ativo'):
        qtd = float(group['Quantidade'].sum())
        if qtd > 0:
            pm = (group['Quantidade'] * group['Preço Médio']).sum() / qtd
            ts = sum((pd.Timestamp(r['Data Média']).timestamp() * r['Quantidade']) for _, r in group.iterrows() if pd.notna(r['Data Média']))
            data_media = pd.to_datetime(ts/qtd, unit='s').date()
            linhas.append({
                "Ativo": ativo, 
                "Quantidade": qtd, 
                "Preço Médio": float(pm), 
                "Data Média": data_media
            })
            
    return pd.DataFrame(linhas)

def corrigir_cabecalho_b3(df):
    if df.empty: 
        return df
    if 'Data do Negócio' in df.columns or 'Data Média' in df.columns: 
        return df
        
    def tentar_mapear(lista_colunas):
        lista_upper = [str(c).strip().upper() for c in lista_colunas]
        is_negociacao = any('DATA DO' in c or 'NEGÓCIO' in c or 'NEGOCIO' in c for c in lista_upper)
        is_mov = any('PRODUTO' in c or 'ATIVO' in c for c in lista_upper) and any('MOVIMENTA' in c or 'MOVIMEN' in c for c in lista_upper)
        
        if is_negociacao or is_mov:
            mapeamento = {}
            for col in lista_colunas:
                c_up = str(col).strip().upper()
                if 'DAT' in c_up: mapeamento[col] = 'Data do Negócio'
                elif 'PROD' in c_up or 'ATIV' in c_up or 'CÓD' in c_up or 'COD' in c_up or 'PAPEL' in c_up: mapeamento[col] = 'Código de Negociação'
                elif 'MOV' in c_up or 'TIP' in c_up: mapeamento[col] = 'Tipo de Movimentação'
                elif 'VAL' in c_up: mapeamento[col] = 'Valor'
                elif 'QUA' in c_up or 'QTD' in c_up: mapeamento[col] = 'Quantidade'
                elif 'PRE' in c_up or 'AJUSTE' in c_up: mapeamento[col] = 'Preço Unitário'
                elif 'ENTRADA' in c_up or 'SAÍDA' in c_up or 'SAIDA' in c_up or 'C/D' in c_up: mapeamento[col] = 'Entrada/Saída'
            return mapeamento, is_mov
        return None, False

    mapeamento, is_mov = tentar_mapear(df.columns.tolist())
    
    if mapeamento and ('Data do Negócio' in mapeamento.values() or 'Código de Negociação' in mapeamento.values()):
        df = df.rename(columns=mapeamento)
        if is_mov and 'Valor' not in df.columns and 'Preço Unitário' in df.columns and 'Quantidade' in df.columns:
            df['Valor'] = df['Quantidade'].apply(limpar_numero) * df['Preço Unitário'].apply(limpar_numero)
        return df

    for i in range(min(30, len(df))):
        linha = df.iloc[i].tolist()
        mapeamento, is_mov = tentar_mapear(linha)
        if mapeamento:
            df.columns = linha
            df_sub = df.iloc[i+1:].reset_index(drop=True).rename(columns=mapeamento)
            if is_mov and 'Valor' not in df_sub.columns and 'Preço Unitário' in df_sub.columns and 'Quantidade' in df_sub.columns:
                df_sub['Valor'] = df_sub['Quantidade'].apply(limpar_numero) * df_sub['Preço Unitário'].apply(limpar_numero)
            return df_sub
            
    return df

def processar_planilha_b3(df):
    if df.empty: 
        return pd.DataFrame()
        
    df['Data do Negócio'] = pd.to_datetime(df['Data do Negócio'], dayfirst=True, errors='coerce')
    df['Quantidade'] = df['Quantidade'].apply(limpar_numero)
    
    if 'Valor' in df.columns: 
        df['Valor'] = df['Valor'].apply(limpar_numero)
    elif 'Preço Unitário' in df.columns: 
        df['Valor'] = df['Quantidade'] * df['Preço Unitário'].apply(limpar_numero)
    else: 
        df['Valor'] = 0.0

    df = df.sort_values('Data do Negócio')
    posicoes = {}
    
    for _, row in df.iterrows():
        if pd.isna(row.get('Código de Negociação')): 
            continue
            
        ticker_raw = str(row['Código de Negociação']).strip().upper()
        if " - " in ticker_raw: 
            ticker_raw = ticker_raw.split(" - ")[0].strip()
        elif " " in ticker_raw: 
            ticker_raw = ticker_raw.split(" ")[0].strip()
            
        ticker = MAPEAMENTO_TICKERS.get(ticker_raw[:-1] if ticker_raw.endswith('F') and len(ticker_raw) > 4 else ticker_raw, ticker_raw[:-1] if ticker_raw.endswith('F') and len(ticker_raw) > 4 else ticker_raw)
        
        if not any(c.isdigit() for c in ticker) or len(ticker) < 4: 
            continue
            
        if 'Tipo de Movimentação' in df.columns:
            mov_tipo_str = str(row['Tipo de Movimentação']).strip().upper()
            if any(p in mov_tipo_str for p in ['RENDIMENTO', 'JUROS', 'DIVIDENDO', 'JCP', 'REEMBOLSO', 'BONIFICAÇÃO', 'DESDOBRAMENTO', 'GRUPAMENTO', 'FRAÇÃO']):
                continue
        
        qtd = row['Quantidade']
        valor = row['Valor']
        data = row['Data do Negócio'] if pd.notna(row['Data do Negócio']) else pd.Timestamp.now()
        
        if ticker not in posicoes: 
            posicoes[ticker] = {'qtd': 0.0, 'valor': 0.0, 'ts_medio': 0.0}
        
        is_compra = False
        is_venda = False
        
        if 'Entrada/Saída' in df.columns and pd.notna(row['Entrada/Saída']):
            io_dir = str(row['Entrada/Saída']).strip().upper()
            if 'CRED' in io_dir or 'ENT' in io_dir or 'C' == io_dir: 
                is_compra = True
            elif 'DEB' in io_dir or 'SAI' in io_dir or 'D' == io_dir: 
                is_venda = True
        
        if not is_compra and not is_venda and 'Tipo de Movimentação' in df.columns:
            tipo_mov = str(row['Tipo de Movimentação']).strip().upper()
            if any(term in tipo_mov for term in ['COMPRA', 'CREDITO', 'CRÉDITO', 'TRANSFERÊNCIA', 'TRANSFERENCIA']) or tipo_mov == 'C':
                is_compra = True
            elif any(term in tipo_mov for term in ['VENDA', 'DEBITO', 'DÉBITO']) or tipo_mov == 'V':
                is_venda = True
        
        if is_compra:
            q_ant = posicoes[ticker]['qtd']
            ts_ant = posicoes[ticker]['ts_medio']
            ts_novo = pd.Timestamp(data).timestamp()
            
            posicoes[ticker]['ts_medio'] = ts_novo if q_ant == 0 else ((ts_ant * q_ant) + (ts_novo * qtd)) / (q_ant + qtd)
            posicoes[ticker]['qtd'] += qtd
            posicoes[ticker]['valor'] += valor
            
        elif is_venda:
            if qtd >= (posicoes[ticker]['qtd'] - 0.001): 
                posicoes[ticker] = {'qtd': 0.0, 'valor': 0.0, 'ts_medio': 0.0}
            else:
                pm = posicoes[ticker]['valor'] / posicoes[ticker]['qtd'] if posicoes[ticker]['qtd'] > 0 else 0
                posicoes[ticker]['qtd'] -= qtd
                posicoes[ticker]['valor'] -= (qtd * pm)
                
    ativos_finais = []
    for k, v in posicoes.items():
        if v['qtd'] > 0:
            ativos_finais.append({
                "Ativo": k, 
                "Quantidade": v['qtd'], 
                "Preço Médio": v['valor']/v['qtd'], 
                "Data Média": pd.to_datetime(v['ts_medio'], unit='s').date()
            })
            
    return consolidar_carteira(pd.DataFrame(ativos_finais))

# ==========================================
# 5. SIDEBAR: UPLOAD, INTEGRAÇÃO B3 E BACKUP
# ==========================================
with st.sidebar:
    st.markdown("### 👤 ANALISTA OPERACIONAL")
    if st.button("🚪 Sair do Terminal", use_container_width=True):
        st.session_state.clear()
        st.rerun()
        
    st.divider()
    
    st.markdown("### 💾 Segurança em Nuvem")
    if st.button("Sincronizar no Banco de Dados", type="primary", use_container_width=True):
        salvar_dados_completos_db(st.session_state.username)
        st.success("Dados blindados no banco com sucesso!")

    # --- NOVO: PAINEL ADMIN PROTEGIDO (Só aparece para Rodrigo) ---
    if st.session_state.username.strip().lower() == "rodrigo":
        st.divider()
        st.markdown("### 👑 Painel Admin")
        if st.button("Ver Relatório de Usuários", use_container_width=True):
            conn = get_db_connection()
            c = conn.cursor()
            c.execute("SELECT username FROM usuarios")
            usuarios_db = c.fetchall()
            conn.close()
            st.info(f"**Total de Contas Registradas:** {len(usuarios_db)}")
            for u in usuarios_db:
                st.write(f"👤 {u[0]}")
        
    st.divider()
    
    st.markdown("### 1. Integrar Notas B3 (Opcional)")
    st.info("Faça o upload de planilhas para substituir toda a base ou apenas integrar novas operações à carteira atual.")
    
    arquivo_principal = st.file_uploader("Substituir Base Completa (B3 / CSV Backup)", type=["xlsx", "csv"])
    arquivo_novo = st.file_uploader("Integrar Novas Operações", type=["xlsx", "csv"])
    
    if arquivo_novo:
        data_corte = st.date_input("Filtrar Novas a partir de:", pd.Timestamp.now().date() - pd.Timedelta(days=15))
    else:
        data_corte = None

    if st.button("🚀 Processar Arquivos B3", use_container_width=True):
        base_atual = st.session_state.df_base.copy()
        
        if arquivo_principal:
            if arquivo_principal.name.endswith('.csv'):
                txt = arquivo_principal.getvalue().decode('utf-8-sig', errors='ignore')
                sep_detectado = '\t' if txt.count('\t') > txt.count(';') else ';'
                df_p = pd.read_csv(io.StringIO(txt), sep=sep_detectado)
            else:
                df_p = pd.read_excel(arquivo_principal)
                
            df_p = corrigir_cabecalho_b3(df_p)
            
            if 'Data Média' in df_p.columns: 
                base_atual = consolidar_carteira(df_p)
            elif 'Data do Negócio' in df_p.columns: 
                base_atual = processar_planilha_b3(df_p)
            else: 
                st.error("Formato inválido. Planilha B3 não reconhecida.")
                st.stop()
                
        if arquivo_novo and not base_atual.empty:
            if arquivo_novo.name.endswith('.csv'):
                txt_n = arquivo_novo.getvalue().decode('utf-8-sig', errors='ignore')
                sep_detectado_n = '\t' if txt_n.count('\t') > txt_n.count(';') else ';'
                df_n = pd.read_csv(io.StringIO(txt_n), sep=sep_detectado_n)
            else:
                df_n = pd.read_excel(arquivo_novo)
                
            df_n = corrigir_cabecalho_b3(df_n)
            
            if not df_n.empty and 'Data do Negócio' in df_n.columns:
                df_n['Data do Negócio'] = pd.to_datetime(df_n['Data do Negócio'], dayfirst=True, errors='coerce')
                df_n = df_n[df_n['Data do Negócio'].dt.date >= data_corte]
                
                linhas_b = []
                for _, r in base_atual.iterrows():
                    linhas_b.append({
                        "Código de Negociação": r['Ativo'], 
                        "Tipo de Movimentação": "Compra", 
                        "Data do Negócio": pd.to_datetime(r['Data Média']), 
                        "Quantidade": r['Quantidade'], 
                        "Valor": r['Quantidade'] * r['Preço Médio']
                    })
                    
                base_atual = processar_planilha_b3(pd.concat([pd.DataFrame(linhas_b), df_n], ignore_index=True))
                
        st.session_state.df_base = base_atual
        st.warning("Memória atualizada. Lembre-se de salvar no DB.")
        st.rerun()
# ==========================================
# 6. PAINEL MACRO E CONTROLE MANUAL DE ATIVOS
# ==========================================
st.markdown("### 👑 Conjuntura Macroeconômica")
c_m1, c_m2 = st.columns([1, 2])

c_m1.success(
    f"🎯 **Cenário Atual (Vigente)**\n\n"
    f"Selic Vigente: **{f_pct(selic_hoje)} a.a.**\n\n"
    f"IPCA 12 meses: **{f_pct(ipca_12m_hoje)}**"
)

c_m2.info(
    f"🔮 **Projeções do Mercado (Boletim Focus)**\n\n"
    f"**Taxa Selic:** {ano_atual}: **{f_pct(proj_focus.get(f'Selic_{ano_atual}', 0))}** |  "
    f"{ano_atual+1}: **{f_pct(proj_focus.get(f'Selic_{ano_atual+1}', 0))}** |  "
    f"{ano_atual+2}: **{f_pct(proj_focus.get(f'Selic_{ano_atual+2}', 0))}**\n\n"
    f"**Inflação (IPCA):** {ano_atual}: **{f_pct(proj_focus.get(f'IPCA_{ano_atual}', 0))}** |  "
    f"{ano_atual+1}: **{f_pct(proj_focus.get(f'IPCA_{ano_atual+1}', 0))}** |  "
    f"{ano_atual+2}: **{f_pct(proj_focus.get(f'IPCA_{ano_atual+2}', 0))}**"
)

st.write("---")
st.markdown("### 2. Controle Operacional da Carteira")

c_op1, c_op2, c_op3 = st.columns([1, 1.5, 1])

with c_op1:
    ativos_lista = sorted(st.session_state.df_base["Ativo"].tolist()) if not st.session_state.df_base.empty else []
    tdel = st.selectbox("Excluir Ativo da Carteira:", [""] + ativos_lista)
    if st.button("Remover Ativo Selecionado", use_container_width=True) and tdel:
        st.session_state.df_base = st.session_state.df_base[st.session_state.df_base["Ativo"] != tdel]
        st.rerun()

with c_op2:
    nt = st.text_input("Novo Aporte ou Ativo Manual (Ticker)")
    tipo_op = st.selectbox("Tipo de Operação:", ["Compra", "Venda"], key="tipo_op_manual")
    cq, cp, cd = st.columns([1, 1, 1.2])
    nq = cq.number_input("Quantidade", min_value=1)
    np_v = cp.number_input("Preço Unitário (R$)", min_value=0.01)
    nd_v = cd.date_input("Data da Operação", value=pd.Timestamp.now().date())
    
    if st.button("Adicionar à Carteira / Integrar", use_container_width=True) and nt:
        ticker = nt.upper().strip()
        df_base = st.session_state.df_base.copy()
        
        pm_epoca = 0.0
        resultado_realizado = 0.0
        
        # Verifica a posição atual para obter o preço médio histórico antes da mudança
        if not df_base.empty and ticker in df_base['Ativo'].values:
            pm_epoca = float(df_base[df_base['Ativo'] == ticker]['Preço Médio'].values[0])
            qtd_atual = float(df_base[df_base['Ativo'] == ticker]['Quantidade'].values[0])
        else:
            qtd_atual = 0.0
            
        if tipo_op == "Compra":
            nova_linha = pd.DataFrame([{
                "Ativo": ticker, 
                "Quantidade": float(nq), 
                "Preço Médio": float(np_v), 
                "Data Média": nd_v
            }])
            st.session_state.df_base = consolidar_carteira(pd.concat([df_base, nova_linha], ignore_index=True))
            resultado_realizado = 0.0 # Compra não gera lucro/prejuízo realizado imediato
            
        elif tipo_op == "Venda":
            if qtd_atual >= nq:
                # Lucro/Prejuízo = (Preço de Venda - Custo Médio de Aquisição) * Quantidade
                resultado_realizado = (float(np_v) - pm_epoca) * float(nq)
                
                # Deduz a quantidade vendida da carteira ativa
                df_base.loc[df_base['Ativo'] == ticker, 'Quantidade'] -= float(nq)
                
                # Se a posição foi zerada por completo, extingue a linha da carteira ativa
                if df_base.loc[df_base['Ativo'] == ticker, 'Quantidade'].values[0] <= 0:
                    df_base = df_base[df_base['Ativo'] != ticker]
                    
                st.session_state.df_base = df_base
            else:
                st.error(f"Operação Bloqueada: Você possui apenas {int(qtd_atual)} cotas de {ticker}. Não pode vender {nq}.")
                st.stop()
                
        # --- REGISTRAR NO LEDGER DE TRANSAÇÕES PERMANENTE ---
        nova_tx = pd.DataFrame([{
            "Ativo": ticker,
            "Tipo": tipo_op,
            "Quantidade": float(nq),
            "Preço Unitário": float(np_v),
            "Preço Médio na Época": pm_epoca,
            "Resultado Realizado": resultado_realizado,
            "Data": nd_v
        }])
        
        if 'df_transacoes' not in st.session_state or st.session_state.df_transacoes.empty:
            st.session_state.df_transacoes = nova_tx
        else:
            st.session_state.df_transacoes = pd.concat([st.session_state.df_transacoes, nova_tx], ignore_index=True)
            
        if st.session_state.username:
            salvar_dados_completos_db(st.session_state.username)
        st.rerun()

with c_op3:
    if not st.session_state.df_base.empty:
        csv_data = st.session_state.df_base.to_csv(index=False, sep=';', encoding='utf-8-sig')
        st.download_button(
            label="📥 Baixar Carteira Ajustada (CSV Backup)", 
            data=csv_data, 
            file_name=f"Carteira_Ajustada_{st.session_state.username}.csv", 
            use_container_width=True
        )
    st.write("")
    if st.button("💾 Salvar no Banco de Dados", type="primary", use_container_width=True):
        salvar_dados_completos_db(st.session_state.username)
        st.success("Dados blindados no banco com sucesso!")

st.markdown("#### 📝 Tabela Editável (Ajuste Fino e Resgate Histórico)")
st.info("Ajuste as quantidades, os preços médios, ou a **Data de Compra** (o robô usará essa data para resgatar todos os dividendos passados a que você tem direito).")

df_editado = st.data_editor(
    st.session_state.df_base, 
    use_container_width=True, 
    hide_index=True,
    column_config={
        "Data Média": st.column_config.DateColumn("Data de Compra (Início dos Divs)", format="DD/MM/YYYY"),
        "Preço Médio": st.column_config.NumberColumn("Preço Médio (R$)", format="%.2f")
    }
)

if st.button("🚀 Conectar ao Mercado Vivo", type="primary", use_container_width=True):
    if not df_editado.empty:
        st.session_state.df_base = consolidar_carteira(df_editado)
        
        prg = st.progress(0)
        total = len(st.session_state.df_base)
        dm = {}
        sim = []
        l_novos_prov = [] # Inicializa o rastreador para o Livro-Razão
        
        for i, r in st.session_state.df_base.iterrows():
            t = str(r['Ativo']).upper()
            dc = pd.to_datetime(r['Data Média']) if pd.notna(r['Data Média']) else pd.Timestamp.now()
            
            p_at = float(r['Preço Médio'])
            d_tot = 0.0
            d_12m = 0.0
            is_fii = t.endswith('11') and t not in UNITS_ACOES
            setor = "Desconhecido"
            
            try:
                tk = yf.Ticker(f"{t}.SA")
                h = tk.history(period="1d")
                if not h.empty: 
                    p_at = float(h['Close'].iloc[-1])
                    
                dvs = tk.dividends
                if not dvs.empty:
                    if dvs.index.tz is not None: 
                        dvs.index = dvs.index.tz_localize(None)
                    
                    dvs_validos = dvs[dvs.index >= pd.Timestamp(dc)]
                    d_tot = float(dvs_validos.sum() * r['Quantidade'])
                    d_12m = float(dvs[dvs.index >= (pd.Timestamp.now() - pd.DateOffset(years=1))].sum())
                    
                    # --- NOVO: GERANDO ENTRADAS PARA O LIVRO-RAZÃO ---
                    for d_idx, val_h in dvs_validos.items():
                        l_novos_prov.append({
                            "Ativo": t, 
                            "Data Ex": d_idx.date(), 
                            "Valor Recebido (R$)": float(val_h * r['Quantidade'])
                        })
                        
                setor = traduzir_setor(tk.info.get('industry', ''))
            except: 
                pass
            
            cdi_ac, ipca_ac = 0.0, 0.0
            try: 
                if not df_macro.empty:
                    f_m = df_macro.loc[dc:]
                    cdi_ac = ((1 + f_m['CDI'].dropna()).prod() - 1) * 100
                    ipca_ac = ((1 + f_m['IPCA'].dropna()).prod() - 1) * 100
            except: 
                pass
            
            vpa = fundamentos_br.get(t, {}).get('vpa', 0)
            lpa = fundamentos_br.get(t, {}).get('lpa', 0)
            
            dm[t] = {
                "Qtd": float(r['Quantidade']), 
                "PM": float(r['Preço Médio']), 
                "Data": dc, 
                "Preço Atual": p_at, 
                "Div_Total": d_tot, 
                "CDI": cdi_ac, 
                "IPCA": ipca_ac, 
                "Tipo": "FII" if is_fii else "Ação", 
                "Setor": setor
            }
            
            sim.append({
                "Ativo": t, 
                "Cotação Atual": p_at, 
                "VPA (Contábil)": vpa, 
                "LPA Projetado": lpa, 
                "Div. Projetado (R$)": d_12m
            })
            
            prg.progress((i + 1) / total)
            
        # --- NOVO: SALVAMENTO INTELIGENTE NO LIVRO-RAZÃO ---
        if l_novos_prov:
            df_n = pd.DataFrame(l_novos_prov)
            if 'df_ledger' not in st.session_state or st.session_state.df_ledger.empty:
                st.session_state.df_ledger = df_n
            else:
                # Junta o histórico antigo com o novo e remove duplicatas (mantendo a conta mais nova)
                df_c = pd.concat([df_n, st.session_state.df_ledger])
                df_c['Data Ex'] = pd.to_datetime(df_c['Data Ex']).dt.date
                st.session_state.df_ledger = df_c.drop_duplicates(subset=['Ativo', 'Data Ex'], keep='first')
            
        st.session_state.dados_mercado = dm
        st.session_state.df_simul = pd.DataFrame(sim)
        st.success("Mercado Sincronizado com a B3! As abas de análise estão prontas e atualizadas.")
    else: 
        st.warning("A carteira está vazia. Adicione ativos antes de conectar.")

st.write("---")
# ==========================================
# 7. DASHBOARDS E RELATÓRIOS (TABS DE ANÁLISE)
# ==========================================
tab_visao, tab_val, tab_radar, tab_radar_fii, tab_graf, tab_prov, tab_tesouro, tab_extrato, tab_ia = st.tabs([
    "📊 Visão Geral", 
    "💰 Valuation", 
    "🎯 Radar & Projeção",
    "🏢 Radar FIIs", 
    "📈 Gráficos", 
    "💸 Proventos B3", 
    "🏛️ Tesouro Direto", 
    "📜 Extrato & Lucros", 
    "💬 Gestora IA (CNPI)"
])

if st.session_state.dados_mercado:
    l_pf = []
    for t, dm in st.session_state.dados_mercado.items():
        investido = dm['Qtd'] * dm['PM']
        saldo = dm['Qtd'] * dm['Preço Atual']
        div_acumulado = dm['Div_Total']
        
        l_pf.append({
            "Ativo": t, 
            "Tipo": dm["Tipo"], 
            "Setor": dm.get("Setor", "Desconhecido"), 
            "Qtd": int(dm['Qtd']), 
            "Preço Médio": dm['PM'], 
            "Preço Atual": dm['Preço Atual'],
            "Total Investido": investido, 
            "Saldo Atual": saldo, 
            "Saldo C/ Dividendos": saldo + div_acumulado,
            "Resultado (R$)": saldo - investido, 
            "Resultado C/ Dividendos": (saldo - investido) + div_acumulado,
            "Data Média": dm['Data'].strftime('%d/%m/%Y'), 
            "Total Div. (R$)": div_acumulado, 
            "DY on Cost (%)": (div_acumulado / investido) * 100 if investido > 0 else 0, 
            "Evolução c/ Div (%)": (((saldo + div_acumulado) / investido) - 1) * 100 if investido > 0 else 0,
            "IPCA Acum. (%)": dm['IPCA'], 
            "CDI Acum. (%)": dm['CDI']
        })
        
    df_perf_final = pd.DataFrame(l_pf)

    with tab_visao:
        st.markdown("### 🏆 Visão Global da Carteira")
        
        df_a = df_perf_final[df_perf_final['Tipo'] == 'Ação']
        df_f = df_perf_final[df_perf_final['Tipo'] == 'FII']
        
        saldo_acoes = df_a['Saldo Atual'].sum() if not df_a.empty else 0.0
        saldo_fiis = df_f['Saldo Atual'].sum() if not df_f.empty else 0.0
        
        ev_a = (saldo_acoes / df_a['Total Investido'].sum() - 1) * 100 if not df_a.empty and df_a['Total Investido'].sum() > 0 else 0
        ev_f = (saldo_fiis / df_f['Total Investido'].sum() - 1) * 100 if not df_f.empty and df_f['Total Investido'].sum() > 0 else 0

        saldo_tesouro = st.session_state.df_tesouro['Valor Investido (R$)'].apply(limpar_numero).sum() if not st.session_state.df_tesouro.empty else 0.0
        patrimonio_total = saldo_acoes + saldo_fiis + saldo_tesouro

        st.success(f"💰 **Patrimônio Total Consolidado (Ações + FIIs + Tesouro):** {f_brl(patrimonio_total)}")

        m1, m2, m3, m4 = st.columns(4)
        m1.metric("📈 Patrimônio Ações", f_brl(saldo_acoes), f_pct(ev_a))
        m2.metric("🏢 Patrimônio FIIs", f_brl(saldo_fiis), f_pct(ev_f))
        m3.metric("🏛️ Tesouro Direto", f_brl(saldo_tesouro))
        
        tot_divs = st.session_state.df_ledger['Valor Recebido (R$)'].sum() if 'df_ledger' in st.session_state and not st.session_state.df_ledger.empty else df_perf_final['Total Div. (R$)'].sum() if not df_perf_final.empty else 0.0
        m4.metric("💸 Dividendos Históricos (Globais)", f_brl(tot_divs))

        formatacao_t1 = {
            c: f_brl for c in ["Preço Médio", "Preço Atual", "Total Investido", "Saldo Atual", "Saldo C/ Dividendos", "Resultado (R$)", "Resultado C/ Dividendos", "Total Div. (R$)"]
        }
        formatacao_t1.update({
            c: f_pct for c in ["DY on Cost (%)", "Evolução c/ Div (%)", "IPCA Acum. (%)", "CDI Acum. (%)"]
        })
        
        st.dataframe(df_perf_final.drop(columns=['Tipo', 'Setor']).style.format(formatacao_t1), use_container_width=True, hide_index=True)

    with tab_val:
        st.markdown("#### Métodos Certificados de Valuation e Indicadores")
        st.markdown("""
        * **Preço Teto Decio Bazin:** Calcula o preço máximo ideal para compra focado em retornos via dividendos.
        * **Preço Justo Benjamin Graham:** Avalia o valor intrínseco e contábil.
        * **P/L e P/VP:** Indicadores fundamentalistas de preço sobre lucro e patrimônio.
        """)
        
        yd = st.number_input("Taxa de Retorno Mínima Exigida Bazin (%):", value=6.0, step=0.5) / 100.0
        
        df_sim = st.data_editor(st.session_state.df_simul[["Ativo", "Cotação Atual", "Div. Projetado (R$)", "VPA (Contábil)", "LPA Projetado"]], use_container_width=True, hide_index=True, disabled=["Ativo", "Cotação Atual"])
        st.session_state.df_simul[["Div. Projetado (R$)", "VPA (Contábil)", "LPA Projetado"]] = df_sim[["Div. Projetado (R$)", "VPA (Contábil)", "LPA Projetado"]]
        
        rv = []
        for _, r in df_sim.iterrows():
            is_fii = r['Ativo'].endswith('11') and r['Ativo'] not in UNITS_ACOES
            
            bz = (float(r["Div. Projetado (R$)"]) / yd) if float(r["Div. Projetado (R$)"]) > 0 else 0.0
            mbz = ((bz / float(r["Cotação Atual"])) - 1) * 100 if bz > 0 else 0.0
            
            pl_calc = float(r["Cotação Atual"]) / float(r["LPA Projetado"]) if float(r["LPA Projetado"]) > 0 else 0.0
            pvp_calc = float(r["Cotação Atual"]) / float(r["VPA (Contábil)"]) if float(r["VPA (Contábil)"]) > 0 else 0.0
            
            if not is_fii:
                gh = (22.5 * float(r["LPA Projetado"]) * float(r["VPA (Contábil)"])) ** 0.5 if float(r["LPA Projetado"]) > 0 and float(r["VPA (Contábil)"]) > 0 else 0.0
                mgh = ((gh / float(r["Cotação Atual"])) - 1) * 100 if gh > 0 else 0.0
            else:
                gh = np.nan
                mgh = np.nan
                
            rv.append({"Ativo": r['Ativo'], "P/L": pl_calc, "P/VP": pvp_calc, "Teto Bazin": bz, "Margem Bazin (%)": mbz, "Justo Graham": gh, "Margem Graham (%)": mgh})
            
        st.session_state.df_recs_val = pd.DataFrame(rv)
        
        formatacao_t2 = {
            "P/L": lambda x: f"{x:.2f}" if x > 0 else "-",
            "P/VP": lambda x: f"{x:.2f}" if x > 0 else "-",
            "Teto Bazin": lambda x: f_brl(x) if x > 0 else "-", 
            "Justo Graham": lambda x: f_brl(x) if pd.notna(x) and x > 0 else "-", 
            "Margem Bazin (%)": lambda x: f_pct(x) if x != 0 else "-", 
            "Margem Graham (%)": lambda x: f_pct(x) if pd.notna(x) and x != 0 else "-"
        }
        st.dataframe(st.session_state.df_recs_val.style.format(formatacao_t2), use_container_width=True, hide_index=True)

    with tab_radar: 
        st.markdown("##### Parametrização do Radar Operacional")
        c_p1, c_p2, c_p3, c_p4 = st.columns(4)
        patr_fora = c_p1.number_input("Patrimônio Externo / Caixa (R$):", value=0.0, step=1000.0)
        aporte = c_p2.number_input("Aporte Mensal Previsto (R$):", value=2000.0, step=500.0)
        rent = c_p3.number_input("Rentabilidade Mensal Alvo (%):", value=0.8, step=0.1) / 100.0
        cresc_div = c_p4.number_input("Crescimento Anual de Dividendos (%):", value=5.0, step=1.0) / 100.0

        st.markdown("##### 🎯 Triagem Estratégica Corporativa")
        st.info(
            "**Teoria da Margem de Segurança:**\n\n"
            "**Graham (>15% a 20%):** Protege contra quedas de mercado sistêmicas e balanços irreais. Você só entra se a ação estiver muito barata frente ao patrimônio físico dela.\n"
            "**Bazin (>0% a 5%):** Garante que, independentemente do sobe e desce da bolsa, o dinheiro depositado na sua conta via dividendos cobrirá a taxa que você estipulou."
        )
        
        c_m1, c_m2 = st.columns(2)
        mb_ex = c_m1.number_input("Margem Mínima Bazin Exigida (%):", value=5.0)
        mg_ex = c_m2.number_input("Margem Mínima Graham Exigida (%):", value=15.0)
        
        df_radar = pd.merge(df_perf_final[['Ativo', 'Tipo', 'Preço Atual']], st.session_state.df_recs_val, on='Ativo')
        
        status_bazin = []
        status_graham = []
        
        for _, row in df_radar.iterrows():
            if row['Teto Bazin'] > 0:
                if row['Margem Bazin (%)'] >= mb_ex: 
                    status_bazin.append("COMPRA 🟢")
                elif row['Margem Bazin (%)'] >= -5: 
                    status_bazin.append("MANTER 🟡")
                else: 
                    status_bazin.append("VENDA 🔴")
            else:
                status_bazin.append("MANTER 🟡")
                
            if row['Tipo'] == 'Ação' and pd.notna(row['Justo Graham']) and row['Justo Graham'] > 0:
                if row['Margem Graham (%)'] >= mg_ex: 
                    status_graham.append("COMPRA 🟢")
                elif row['Margem Graham (%)'] >= 0: 
                    status_graham.append("MANTER 🟡")
                else: 
                    status_graham.append("VENDA 🔴")
            else:
                status_graham.append("-")
                
        df_radar['Status Bazin'] = status_bazin
        df_radar['Status Graham'] = status_graham
        
        st.dataframe(df_radar[['Ativo', 'Tipo', 'Preço Atual', 'Teto Bazin', 'Margem Bazin (%)', 'Status Bazin', 'Justo Graham', 'Margem Graham (%)', 'Status Graham']].style.format(formatacao_t2 | {"Preço Atual": f_brl}), use_container_width=True, hide_index=True)

        st.markdown("##### ❄️ Projeção Bola de Neve (1 Ano)")
        saldo_dinamico = df_perf_final['Saldo Atual'].sum() + patr_fora
        b_div = st.session_state.df_simul['Div. Projetado (R$)'].sum() / 12 if not st.session_state.df_simul.empty else 0.0
        
        ac_ap = 0.0
        ac_jd = 0.0
        lp = []
        
        for m in range(13):
            mes_label = f"Mês {m:02d}" 
            if m == 0:
                lp.append({"Mês": mes_label, "Capital Inicial": saldo_dinamico, "Aportes Acumulados": 0.0, "Juros/Divs Acumulados": 0.0, "Total Projetado": saldo_dinamico})
            else:
                gc = saldo_dinamico * rent
                div_m = b_div * ((1 + cresc_div) ** (m/12))
                ac_jd += (gc + div_m)
                ac_ap += aporte
                saldo_dinamico += (gc + div_m + aporte)
                
                lp.append({
                    "Mês": mes_label, 
                    "Capital Inicial": df_perf_final['Saldo Atual'].sum() + patr_fora, 
                    "Aportes Acumulados": ac_ap, 
                    "Juros/Divs Acumulados": ac_jd,
                    "Total Projetado": saldo_dinamico
                })
                
        df_proj_plot = pd.DataFrame(lp)
        df_melt_proj = df_proj_plot.melt(id_vars=["Mês", "Total Projetado"], value_vars=["Capital Inicial", "Aportes Acumulados", "Juros/Divs Acumulados"], var_name="Componente", value_name="Valor (R$)")
        
        fig_proj = px.bar(df_melt_proj, x="Mês", y="Valor (R$)", color="Componente", title="Evolução Patrimonial Controlada", color_discrete_sequence=['#1f4e78', '#00a896', '#f4a261'])
        
        fig_proj.add_scatter(x=df_proj_plot["Mês"], y=df_proj_plot["Total Projetado"], mode='text', text=df_proj_plot["Total Projetado"].apply(lambda x: f"R$ {x/1000:.1f}k"), textposition='top center', showlegend=False)
        fig_proj.update_layout(barmode='stack', yaxis_title="Patrimônio (R$)", xaxis_title="Evolução Mensal")
        
        st.plotly_chart(fig_proj, use_container_width=True)

    with tab_graf:
        st.markdown("#### 📊 Gráficos de Distribuição Patrimonial")
        
        import plotly.express as px
        import pandas as pd
        import yfinance as yf
        
        df_graf = df_perf_final[df_perf_final['Saldo Atual'] > 0] if ('df_perf_final' in locals() and not df_perf_final.empty) else pd.DataFrame()
        
        if not df_graf.empty:
            cg1, cg2 = st.columns(2)
            paleta = ['#003f5c', '#2f4b7c', '#665191', '#a05195', '#d45087', '#f95d6a', '#ff7c43', '#ffa600']
            
            fig_g1 = px.pie(df_graf, values='Saldo Atual', names='Ativo', title="Por Ativo", color_discrete_sequence=paleta)
            cg1.plotly_chart(fig_g1, use_container_width=True, key="pizza_ativos_unica")
            
            fig_g2 = px.pie(df_graf, values='Saldo Atual', names='Tipo', title="Por Classe Operacional", color_discrete_sequence=['#1f4e78', '#00a896'])
            cg2.plotly_chart(fig_g2, use_container_width=True, key="pizza_classes_unica")
        else:
            st.warning("Conecte ao mercado ou adicione ativos com saldo para gerar os gráficos de pizza.")
            
        st.markdown("---")
        st.markdown("#### 📈 Comparativo Histórico e Indexadores")
        
        if not df_graf.empty:
            ativos_disponiveis = sorted(df_graf['Ativo'].unique().tolist())
            c_f_g1, c_f_g2 = st.columns(2)
            
            atv_sel = c_f_g1.multiselect("Comparar Ativos Específicos:", options=ativos_disponiveis, default=ativos_disponiveis[:5] if len(ativos_disponiveis) >= 5 else ativos_disponiveis, key="ms_graf_atv")
            ind_sel = c_f_g2.multiselect("Comparar com os Indexadores:", ['CDI', 'IPCA'], default=['CDI', 'IPCA'], key="ms_graf_ind")
            
            janela = st.radio("Período de Análise:", ["Desde a Data de Compra (Automático)", "Definir Período Customizado (Manual)"], horizontal=True, key="rd_graf_periodo")
            
            if janela == "Desde a Data de Compra (Automático)":
                if atv_sel:
                    df_comp = df_graf[df_graf['Ativo'].isin(atv_sel)][['Ativo', 'Evolução c/ Div (%)', 'CDI Acum. (%)', 'IPCA Acum. (%)']].copy()
                    df_comp = df_comp.rename(columns={'Evolução c/ Div (%)': 'Carteira (c/ Div)', 'CDI Acum. (%)': 'CDI', 'IPCA Acum. (%)': 'IPCA'})
                    
                    colunas_manter = ['Ativo', 'Carteira (c/ Div)']
                    if 'CDI' in ind_sel: colunas_manter.append('CDI')
                    if 'IPCA' in ind_sel: colunas_manter.append('IPCA')
                    
                    colunas_manter = [c for c in colunas_manter if c in df_comp.columns]
                    df_comp = df_comp[colunas_manter]
                    
                    df_melt = df_comp.melt(id_vars='Ativo', var_name='Indicador', value_name='Rentabilidade (%)')
                    
                    fig_comp = px.bar(
                        df_melt, x='Ativo', y='Rentabilidade (%)', color='Indicador', barmode='group',
                        color_discrete_map={'Carteira (c/ Div)': '#003f5c', 'CDI': '#00a896', 'IPCA': '#f4a261'},
                        title="Rentabilidade Acumulada no Tempo de Posse"
                    )
                    st.plotly_chart(fig_comp, use_container_width=True, key="bar_graf_auto_unica")
                else:
                    st.warning("Selecione ao menos um ativo para visualizar o gráfico.")
            else:
                c_dt1, c_dt2 = st.columns(2)
                dt_ini = c_dt1.date_input("De:", pd.Timestamp.now().date() - pd.Timedelta(days=365), key="dt_graf_ini")
                dt_fim = c_dt2.date_input("Até:", pd.Timestamp.now().date(), key="dt_graf_fim")
                
                if st.button("Gerar Gráfico Comparativo", use_container_width=True, key="btn_graf_comp"):
                    if atv_sel:
                        with st.spinner("Calculando série histórica..."):
                            cdi_m, ipca_m = 0.0, 0.0
                            
                            if 'df_macro' in locals() and not df_macro.empty:
                                try:
                                    f_m = df_macro.loc[dt_ini:dt_fim]
                                    cdi_m = ((1 + f_m['CDI'].dropna()).prod() - 1) * 100
                                    ipca_m = ((1 + f_m['IPCA'].dropna()).prod() - 1) * 100
                                except: pass
                                
                            l_res = []
                            for t in atv_sel:
                                r_atv = 0.0
                                try:
                                    ht = yf.Ticker(f"{t}.SA").history(start=dt_ini, end=dt_fim)
                                    if not ht.empty and len(ht) >= 2:
                                        p_i = float(ht['Close'].iloc[0])
                                        p_f = float(ht['Close'].iloc[-1])
                                        d_p = 0.0
                                        try:
                                            al_d = yf.Ticker(f"{t}.SA").dividends
                                            if not al_d.empty:
                                                if al_d.index.tz is not None: al_d.index = al_d.index.tz_localize(None)
                                                d_p = float(al_d[(al_d.index >= pd.Timestamp(dt_ini)) & (al_d.index <= pd.Timestamp(dt_fim))].sum())
                                        except: pass
                                        r_atv = ((p_f + d_p) / p_i - 1) * 100
                                except: pass
                                
                                it_m = {'Ativo': t, 'Carteira (c/ Div)': r_atv}
                                if 'CDI' in ind_sel: it_m['CDI'] = cdi_m
                                if 'IPCA' in ind_sel: it_m['IPCA'] = ipca_m
                                l_res.append(it_m)
                            
                            df_m_plot = pd.DataFrame(l_res)
                            if not df_m_plot.empty:
                                df_melt_m = df_m_plot.melt(id_vars='Ativo', var_name='Indicador', value_name='Rentabilidade (%)')
                                fig_comp_m = px.bar(
                                    df_melt_m, x='Ativo', y='Rentabilidade (%)', color='Indicador', barmode='group',
                                    color_discrete_map={'Carteira (c/ Div)': '#003f5c', 'CDI': '#00a896', 'IPCA': '#f4a261'},
                                    title=f"Desempenho Customizado de {dt_ini.strftime('%d/%m/%Y')} até {dt_fim.strftime('%d/%m/%Y')}"
                                )
                                st.plotly_chart(fig_comp_m, use_container_width=True, key="bar_graf_manual_unica")
                    else:
                        st.warning("Selecione ao menos um ativo para visualizar o gráfico.")

    with tab_radar_fii:
        st.markdown("### 🏢 Radar Analítico de FIIs")
        st.info("Acompanhe os fundos da sua carteira e analise novas oportunidades. **Clique no cabeçalho das colunas para ordenar.** (Dados em tempo real via B3 e Fundamentus).")
        
        if 'fiis_pesquisados' not in st.session_state:
            st.session_state.fiis_pesquisados = []
        if 'fii_pvp_manual' not in st.session_state:
            st.session_state.fii_pvp_manual = {}
            
        c_fii1, c_fii2, c_fii3 = st.columns([2, 1, 1])
        novo_fii = c_fii1.text_input("Simular FII fora da carteira (Ex: HGLG11):", key="txt_busca_fii")
        
        if c_fii2.button("🔍 Buscar Ativo", use_container_width=True):
            if novo_fii:
                tk_clean = novo_fii.upper().strip()
                if not tk_clean.endswith('11') and len(tk_clean) == 4: 
                    tk_clean += '11'
                if tk_clean not in st.session_state.fiis_pesquisados:
                    st.session_state.fiis_pesquisados.append(tk_clean)
                st.rerun()
                    
        if c_fii3.button("🗑️ Limpar Pesquisas", use_container_width=True):
            st.session_state.fiis_pesquisados = []
            st.rerun()
            
        fiis_carteira = df_perf_final[df_perf_final['Tipo'] == 'FII']['Ativo'].tolist() if ('df_perf_final' in locals() and not df_perf_final.empty) else []
        todos_fiis = list(set(fiis_carteira + st.session_state.fiis_pesquisados))
        
        if todos_fiis:
            with st.spinner("Conectando aos servidores do Fundamentus e B3..."):
                dados_fii = []
                import urllib.request
                import re
                
                for f in todos_fiis:
                    pm = df_perf_final[df_perf_final['Ativo'] == f]['Preço Médio'].values[0] if f in fiis_carteira else 0.0
                    
                    preco_mercado = 0.0
                    dy_anual = 0.0
                    dy_mensal = 0.0
                    
                    # 1. Busca Preço e Dividendos (Yahoo Finance Histórico - Imune a bloqueios)
                    try:
                        t = yf.Ticker(f"{f}.SA")
                        hist = t.history(period="1mo")
                        if not hist.empty:
                            preco_mercado = float(hist['Close'].iloc[-1])
                            
                        divs = t.dividends
                        if not divs.empty:
                            if divs.index.tz is not None: 
                                divs.index = divs.index.tz_localize(None)
                            hoje_fii = pd.Timestamp.now()
                            divs_12m = divs[divs.index >= (hoje_fii - pd.DateOffset(months=12))]
                            
                            if preco_mercado > 0:
                                dy_anual = (float(divs_12m.sum()) / preco_mercado) * 100
                                dy_mensal = (float(divs.iloc[-1]) / preco_mercado) * 100
                    except:
                        pass
                        
                    pvp = st.session_state.fii_pvp_manual.get(f, 0.0)
                    
                    # 2. Motor Anti-Cloudflare: Extração de P/VP pelo site do Fundamentus
                    if pvp == 0.0:
                        try:
                            url = f"https://www.fundamentus.com.br/detalhes.php?papel={f}"
                            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'})
                            # Fundamentus usa a codificação iso-8859-1 (padrão brasileiro antigo)
                            html = urllib.request.urlopen(req, timeout=4).read().decode('iso-8859-1')
                            
                            match_pvp = re.search(r'P/VP.*?<span class="txt">([\d,\.]+)</span>', html, re.IGNORECASE | re.DOTALL)
                            if match_pvp:
                                pvp_str = match_pvp.group(1).replace('.', '').replace(',', '.')
                                pvp = float(pvp_str)
                        except: 
                            pass
                    
                    # 3. Fallback final pela sua tabela de Valuation
                    if pvp == 0.0:
                        if 'df_simul' in st.session_state and not st.session_state.df_simul.empty and f in st.session_state.df_simul['Ativo'].values:
                            try:
                                vpa_simul = float(st.session_state.df_simul.loc[st.session_state.df_simul['Ativo'] == f, 'VPA (Contábil)'].values[0])
                                if vpa_simul > 0 and preco_mercado > 0: 
                                    pvp = preco_mercado / vpa_simul
                            except: 
                                pass
                    
                    var_pm = ((preco_mercado / pm) - 1) * 100 if pm > 0 and preco_mercado > 0 else 0.0
                    
                    dados_fii.append({
                        "Ativo": f,
                        "Origem": "Na Carteira" if f in fiis_carteira else "Pesquisado",
                        "Preço Mercado": preco_mercado,
                        "Preço Médio": pm,
                        "Var. Mercado / PM (%)": var_pm,
                        "P/VP": float(pvp) if pd.notna(pvp) else 0.0,
                        "DY Mensal (%)": dy_mensal,
                        "DY Anual (%)": dy_anual
                    })
                    
                if dados_fii:
                    df_radar_fii = pd.DataFrame(dados_fii)
                    
                    df_editado_fii = st.data_editor(
                        df_radar_fii,
                        key="editor_radar_fiis_unique",
                        use_container_width=True,
                        hide_index=True,
                        disabled=["Ativo", "Origem", "Preço Mercado", "Preço Médio", "Var. Mercado / PM (%)", "DY Mensal (%)", "DY Anual (%)"],
                        column_config={
                            "Preço Mercado": st.column_config.NumberColumn("Preço Mercado", format="R$ %.2f"),
                            "Preço Médio": st.column_config.NumberColumn("Preço Médio", format="R$ %.2f"),
                            "Var. Mercado / PM (%)": st.column_config.NumberColumn("Var. Mercado / PM", format="%.2f%%"),
                            "P/VP": st.column_config.NumberColumn("P/VP", format="%.2f", min_value=0.0, max_value=5.0, step=0.01),
                            "DY Mensal (%)": st.column_config.NumberColumn("DY Mensal", format="%.2f%%"),
                            "DY Anual (%)": st.column_config.NumberColumn("DY Anual", format="%.2f%%"),
                        }
                    )
                    
                    if st.button("💾 Salvar Ajustes Manuais de P/VP", use_container_width=True):
                        for _, row in df_editado_fii.iterrows():
                            st.session_state.fii_pvp_manual[row['Ativo']] = float(row['P/VP'])
                        st.success("Fundamentos atualizados e salvos com sucesso!")
                        st.rerun()
        else:
            st.info("Nenhum FII identificado. Comece inserindo um ativo no campo acima.")
        cg1, cg2 = st.columns(2)
        
        paleta = ['#003f5c', '#2f4b7c', '#665191', '#a05195', '#d45087', '#f95d6a', '#ff7c43', '#ffa600']
        
        cg1.plotly_chart(px.pie(df_perf_final, values='Saldo Atual', names='Ativo', title="Por Ativo", color_discrete_sequence=paleta), use_container_width=True)
        cg2.plotly_chart(px.pie(df_perf_final, values='Saldo Atual', names='Tipo', title="Por Classe Operacional", color_discrete_sequence=['#1f4e78', '#00a896']), use_container_width=True)
        
        st.markdown("---")
        st.markdown("#### 📊 Comparativo Histórico e Indexadores")
        
        ativos_disponiveis = sorted(df_perf_final['Ativo'].unique().tolist())
        c_f_g1, c_f_g2 = st.columns(2)
        
        atv_sel = c_f_g1.multiselect("Comparar Ativos Específicos:", options=ativos_disponiveis, default=ativos_disponiveis[:5] if len(ativos_disponiveis) >= 5 else ativos_disponiveis)
        ind_sel = c_f_g2.multiselect("Comparar com os Indexadores:", ['CDI', 'IPCA'], default=['CDI', 'IPCA'])
        
        janela = st.radio("Período de Análise:", ["Desde a Data de Compra (Automático)", "Definir Período Customizado (Manual)"], horizontal=True)
        
        if janela == "Desde a Data de Compra (Automático)":
            if atv_sel:
                df_comp = df_perf_final[df_perf_final['Ativo'].isin(atv_sel)][['Ativo', 'Evolução c/ Div (%)', 'CDI Acum. (%)', 'IPCA Acum. (%)']].copy()
                df_comp = df_comp.rename(columns={'Evolução c/ Div (%)': 'Carteira (c/ Div)', 'CDI Acum. (%)': 'CDI', 'IPCA Acum. (%)': 'IPCA'})
                
                colunas_manter = ['Ativo', 'Carteira (c/ Div)'] + [ind for ind in ind_sel]
                df_comp = df_comp[colunas_manter]
                df_melt = df_comp.melt(id_vars='Ativo', var_name='Indicador', value_name='Rentabilidade (%)')
                
                fig_comp = px.bar(
                    df_melt, x='Ativo', y='Rentabilidade (%)', color='Indicador', barmode='group',
                    color_discrete_map={'Carteira (c/ Div)': '#003f5c', 'CDI': '#00a896', 'IPCA': '#f4a261'},
                    title="Rentabilidade Acumulada no Tempo de Posse"
                )
                st.plotly_chart(fig_comp, use_container_width=True)
            else:
                st.warning("Selecione ao menos um ativo para visualizar o gráfico.")
        else:
            c_dt1, c_dt2 = st.columns(2)
            dt_ini = c_dt1.date_input("De:", pd.Timestamp.now().date() - pd.Timedelta(days=365))
            dt_fim = c_dt2.date_input("Até:", pd.Timestamp.now().date())
            
            if st.button("Gerar Gráfico Comparativo", use_container_width=True):
                if atv_sel:
                    with st.spinner("Calculando série histórica..."):
                        cdi_m, ipca_m = 0.0, 0.0
                        if not df_macro.empty:
                            try:
                                f_m = df_macro.loc[dt_ini:dt_fim]
                                cdi_m = ((1 + f_m['CDI'].dropna()).prod() - 1) * 100
                                ipca_m = ((1 + f_m['IPCA'].dropna()).prod() - 1) * 100
                            except: pass
                            
                        l_res = []
                        for t in atv_sel:
                            r_atv = 0.0
                            try:
                                ht = yf.Ticker(f"{t}.SA").history(start=dt_ini, end=dt_fim)
                                if not ht.empty and len(ht) >= 2:
                                    p_i = ht['Close'].iloc[0]
                                    p_f = ht['Close'].iloc[-1]
                                    d_p = 0.0
                                    try:
                                        al_d = yf.Ticker(f"{t}.SA").dividends
                                        if not al_d.empty:
                                            if al_d.index.tz is not None: al_d.index = al_d.index.tz_localize(None)
                                            d_p = al_d[(al_d.index >= pd.Timestamp(dt_ini)) & (al_d.index <= pd.Timestamp(dt_fim))].sum()
                                    except: pass
                                    r_atv = ((p_f + d_p) / p_i - 1) * 100
                            except: pass
                            
                            it_m = {'Ativo': t, 'Carteira (c/ Div)': r_atv}
                            if 'CDI' in ind_sel: it_m['CDI'] = cdi_m
                            if 'IPCA' in ind_sel: it_m['IPCA'] = ipca_m
                            l_res.append(it_m)
                        
                        df_m_plot = pd.DataFrame(l_res)
                        if not df_m_plot.empty:
                            df_melt_m = df_m_plot.melt(id_vars='Ativo', var_name='Indicador', value_name='Rentabilidade (%)')
                            fig_comp_m = px.bar(
                                df_melt_m, x='Ativo', y='Rentabilidade (%)', color='Indicador', barmode='group',
                                color_discrete_map={'Carteira (c/ Div)': '#003f5c', 'CDI': '#00a896', 'IPCA': '#f4a261'},
                                title=f"Desempenho Customizado de {dt_ini.strftime('%d/%m/%Y')} até {dt_fim.strftime('%d/%m/%Y')}"
                            )
                            st.plotly_chart(fig_comp_m, use_container_width=True)
                else:
                    st.warning("Selecione ao menos um ativo para visualizar o gráfico.")

    with tab_prov:
        st.markdown("### 💸 Proventos Mensais e Status de Pagamento B3")
        cf1, cf2, c_btn = st.columns([2, 2, 2])
        m_map = {1:"Janeiro", 2:"Fevereiro", 3:"Março", 4:"Abril", 5:"Maio", 6:"Junho", 7:"Julho", 8:"Agosto", 9:"Setembro", 10:"Outubro", 11:"Novembro", 12:"Dezembro"}
        
        m_hoje = pd.Timestamp.now().month
        a_hoje = pd.Timestamp.now().year
        
        m_sel = cf1.selectbox("Mês do Provento:", options=list(m_map.keys()), format_func=lambda x: m_map[x], index=m_hoje-1)
        a_sel = cf2.selectbox("Ano de Referência:", options=[a_hoje, a_hoje-1, a_hoje-2])
        
        if c_btn.button("🔄 Processar Renda Mensal", use_container_width=True):
            with st.spinner("Lendo base de dados local (Livro-Razão)..."):
                la, lf = [], []
                
                # --- NOVO: Mapeia Ativos Atuais + Ativos Vendidos que pagaram no mês ---
                ativos_para_processar = set(st.session_state.dados_mercado.keys()) if st.session_state.dados_mercado else set()
                
                df_l = pd.DataFrame()
                if 'df_ledger' in st.session_state and not st.session_state.df_ledger.empty:
                    df_l = st.session_state.df_ledger
                    df_mes_atual = df_l[(pd.to_datetime(df_l['Data Ex']).dt.month == m_sel) & (pd.to_datetime(df_l['Data Ex']).dt.year == a_sel)]
                    ativos_para_processar.update(df_mes_atual['Ativo'].unique())
                
                for t_tk in ativos_para_processar:
                    dm = st.session_state.dados_mercado.get(t_tk, None) if st.session_state.dados_mercado else None
                    
                    val_recebido = 0.0
                    val_unitario = 0.0
                    
                    if not df_l.empty:
                        df_f_mes = df_l[(df_l['Ativo'] == t_tk) & (pd.to_datetime(df_l['Data Ex']).dt.month == m_sel) & (pd.to_datetime(df_l['Data Ex']).dt.year == a_sel)]
                        if not df_f_mes.empty:
                            val_recebido = float(df_f_mes['Valor Recebido (R$)'].sum())
                    
                    if dm:
                        # O ativo ainda está na sua carteira
                        val_unitario = val_recebido / dm['Qtd'] if dm['Qtd'] > 0 else 0.0
                        yoc = (val_recebido / (dm['Qtd'] * dm['PM'])) * 100 if dm['PM'] > 0 else 0
                        dy_m = (val_unitario / dm['Preço Atual']) * 100 if dm['Preço Atual'] > 0 else 0
                        tipo = dm['Tipo']
                        nome_exib = t_tk
                    else:
                        # O ativo foi vendido, mas você tem direito ao provento deste mês!
                        yoc = 0.0
                        dy_m = 0.0
                        tipo = 'FII' if t_tk.endswith('11') else 'Ação'
                        nome_exib = f"{t_tk} (Vendido)"
                        
                    if tipo == 'FII':
                        # Só adiciona ativos vendidos se tiver recebido dinheiro. Ativos atuais mostra sempre (Aguardando)
                        if dm or val_recebido > 0:
                            lf.append({
                                "Fundo (FII)": nome_exib, 
                                "Unitário (R$)": val_unitario, 
                                "Recebido (R$)": val_recebido, 
                                "Yield on Cost (%)": yoc, 
                                "DY Atual (%)": dy_m,
                                "Status": "Pago / Provisionado 🟢" if val_recebido > 0 else "Aguardando 🟡"
                            })
                    else:
                        if val_recebido > 0: 
                            la.append({
                                "Ação": nome_exib, 
                                "Unitário (R$)": val_unitario, 
                                "Recebido (R$)": val_recebido, 
                                "Yield on Cost (%)": yoc, 
                                "DY Atual (%)": dy_m,
                                "Status": "Pago / Provisionado 🟢"
                            })
                
                st.session_state.divs_a = pd.DataFrame(la)
                st.session_state.divs_f = pd.DataFrame(lf)
                st.session_state.divs_m = m_sel
                st.session_state.divs_ano = a_sel
        
        tot_mes = 0.0
        
        if 'divs_f' in st.session_state and not st.session_state.divs_f.empty:
            st.markdown("#### 🏢 Status dos Fundos Imobiliários (FIIs)")
            st.dataframe(st.session_state.divs_f.style.format({"Unitário (R$)": f_brl_4, "Recebido (R$)": f_brl, "Yield on Cost (%)": f_pct, "DY Atual (%)": f_pct}), use_container_width=True, hide_index=True)
            tot_mes += st.session_state.divs_f['Recebido (R$)'].sum()
            
        if 'divs_a' in st.session_state and not st.session_state.divs_a.empty:
            st.markdown("#### 📈 Ações Pagadoras")
            st.dataframe(st.session_state.divs_a.style.format({"Unitário (R$)": f_brl_4, "Recebido (R$)": f_brl, "Yield on Cost (%)": f_pct, "DY Atual (%)": f_pct}), use_container_width=True, hide_index=True)
            tot_mes += st.session_state.divs_a['Recebido (R$)'].sum()
            
        if ('divs_a' in st.session_state and not st.session_state.divs_a.empty) or ('divs_f' in st.session_state and not st.session_state.divs_f.empty):
            st.success(f"**💰 Total Estimado de Proventos no Período ({m_map[st.session_state.divs_m]}/{st.session_state.divs_ano}):** {f_brl(tot_mes)}")
            
        st.markdown("---")
        st.markdown("### 🏛️ Livro-Razão de Proventos (Histórico Permanente)")
        
        if 'df_ledger' in st.session_state and not st.session_state.df_ledger.empty:
            df_hist_tot = st.session_state.df_ledger.sort_values("Data Ex", ascending=False)
            c_h1, c_h2 = st.columns(2)
            atvs_disp = sorted(df_hist_tot['Ativo'].unique().tolist())
            
            atvs_sel = c_h1.multiselect("Filtrar Tabela por Ativo:", options=atvs_disp, default=atvs_disp)
            r_hist = c_h2.date_input("Filtrar Tabela por Período:", value=(min(df_hist_tot['Data Ex']), max(df_hist_tot['Data Ex'])))
            
            df_hist_f = df_hist_tot[df_hist_tot['Ativo'].isin(atvs_sel)]
            if isinstance(r_hist, tuple) and len(r_hist) == 2:
                df_hist_f = df_hist_f[(df_hist_f['Data Ex'] >= r_hist[0]) & (df_hist_f['Data Ex'] <= r_hist[1])]
                
            if not df_hist_f.empty:
                st.dataframe(df_hist_f.style.format({"Valor Recebido (R$)": f_brl}), use_container_width=True, hide_index=True)
                
                xls_buffer = to_excel(df_hist_f, sheet_name="Livro_Razao")
                st.download_button(label="📥 Baixar Livro-Razão (Excel)", data=xls_buffer, file_name=f"Livro_Razao_{st.session_state.username}.xlsx", mime="application/vnd.ms-excel", use_container_width=True)
                
                st.markdown("#### 📈 Parâmetros de Visualização dos Proventos")
                c_gr1, c_gr2, c_gr3 = st.columns([1.5, 1, 1.5])
                
                default_ini = (pd.Timestamp.now() - pd.DateOffset(months=12)).date()
                default_fim = pd.Timestamp.now().date()
                
                r_periodo = c_gr1.date_input("Filtrar Período do Gráfico:", value=(default_ini, default_fim), key="p_graf_div_custom")
                
                atvs_disp_graf = ["Todos os Ativos"] + sorted(df_hist_tot['Ativo'].unique().tolist())
                atv_sel_graf = c_gr2.selectbox("Isolar Ativo no Gráfico:", options=atvs_disp_graf, index=0)
                
                visao_graf = c_gr3.radio("Apresentação Estrutural:", ["Unificado (Consolidado)", "Dividido por Classe (Ação vs FII)"], horizontal=True, key="v_graf_div_custom")
                
                df_chart = df_hist_tot.copy()
                df_chart['Data Ex'] = pd.to_datetime(df_chart['Data Ex']).dt.date
                
                if isinstance(r_periodo, tuple) and len(r_periodo) == 2:
                    df_chart = df_chart[(df_chart['Data Ex'] >= r_periodo[0]) & (df_chart['Data Ex'] <= r_periodo[1])]
                    
                if atv_sel_graf != "Todos os Ativos":
                    df_chart = df_chart[df_chart['Ativo'] == atv_sel_graf]
                    
                if not df_chart.empty:
                    df_chart['Mês/Ano'] = pd.to_datetime(df_chart['Data Ex']).dt.to_period('M').astype(str)
                    df_chart['Tipo'] = df_chart['Ativo'].apply(lambda x: st.session_state.dados_mercado.get(x, {}).get('Tipo', 'FII' if x.endswith('11') else 'Ação'))
                    
                    if visao_graf == "Unificado (Consolidado)":
                        df_chart_g = df_chart.groupby('Mês/Ano')['Valor Recebido (R$)'].sum().reset_index().sort_values('Mês/Ano')
                        fig_divs = px.bar(df_chart_g, x='Mês/Ano', y='Valor Recebido (R$)', text='Valor Recebido (R$)', title="Histórico Consolidado de Proventos", color_discrete_sequence=['#00a896'])
                        fig_divs.update_traces(texttemplate='R$ %{text:,.2f}', textposition='outside')
                    else:
                        df_chart_g = df_chart.groupby(['Mês/Ano', 'Tipo'])['Valor Recebido (R$)'].sum().reset_index().sort_values('Mês/Ano')
                        fig_divs = px.bar(df_chart_g, x='Mês/Ano', y='Valor Recebido (R$)', color='Tipo', barmode='group', title="Histórico de Proventos por Classe (Lado a Lado)", color_discrete_map={'Ação': '#1f4e78', 'FII': '#00a896'})
                        fig_divs.update_traces(texttemplate='R$ %{y:,.2f}', textposition='outside')
                        
                    st.plotly_chart(fig_divs, use_container_width=True)
                else:
                    st.warning("Nenhum registro de provento encontrado para os filtros selecionados no gráfico.")
        else:
            st.info("Nenhum histórico de proventos registrado no Livro-Razão. Conecte ao mercado para iniciar o rastreamento.")

# ==========================================
# 8. ABAS ISOLADAS (SEMPRE ATIVAS)
# ==========================================
with tab_tesouro:
    st.markdown("### 🏛️ Simulador e Controle de Tesouro Direto")
    st.info("Insira os títulos manualmente ou via planilha. O **Valor Atualizado** é calculado dia a dia (pró-rata) com base na taxa desde a data de compra.")
    
    colunas_padrao = ["Título", "Data Compra", "Valor Investido (R$)", "Tipo Taxa", "Taxa Contratada (%)", "Ano Vencimento"]
    if st.session_state.df_tesouro.empty:
        st.session_state.df_tesouro = pd.DataFrame(columns=colunas_padrao)

    with st.expander("➕ Lançar Novo Título Manualmente", expanded=False):
        with st.form("form_add_tesouro", clear_on_submit=True):
            ct1, ct2, ct3 = st.columns(3)
            n_tit = ct1.text_input("Nome do Título", "Tesouro Selic 2029")
            n_tipo = ct2.selectbox("Indexador / Tipo", ["Tesouro Selic", "Pós-fixado (IPCA+)", "Pré-fixado"])
            n_dt = ct3.date_input("Data da Compra", value=pd.Timestamp.now().date())

            ct4, ct5, ct6 = st.columns(3)
            n_inv_str = ct4.text_input("Valor Investido (R$)", value="1000.00")
            n_tx_str = ct5.text_input("Taxa Contratada / Prêmio (%)", value="0.00", help="Use ponto ou vírgula. Ex: 0.15 (Selic) ou 10.5 (Pré-fixado).")
            n_venc_str = ct6.text_input("Ano de Vencimento", value=str(pd.Timestamp.now().year + 3))

            if st.form_submit_button("Registrar Título na Carteira", type="primary"):
                try:
                    n_inv = float(n_inv_str.replace(',', '.'))
                    n_tx = float(n_tx_str.replace(',', '.'))
                    n_venc = int(n_venc_str.strip())
                    
                    novo_t = pd.DataFrame([{
                        "Título": n_tit.strip(), "Data Compra": n_dt, "Valor Investido (R$)": n_inv,
                        "Tipo Taxa": n_tipo, "Taxa Contratada (%)": n_tx, "Ano Vencimento": n_venc
                    }])
                    st.session_state.df_tesouro = pd.concat([st.session_state.df_tesouro, novo_t], ignore_index=True)
                    if st.session_state.username:
                        salvar_dados_completos_db(st.session_state.username)
                    
                    # RERUN REMOVIDO DAQUI. Adicionada apenas a mensagem de sucesso.
                    st.success("Tópico adicionado com sucesso! A tabela abaixo já foi atualizada.")
                except ValueError:
                    st.error("⚠️ Erro de digitação: Certifique-se de inserir apenas números nos campos de Valor, Taxa e Ano.")

    with st.expander("📂 Importar Planilha da B3 / Corretora", expanded=False):
        arq_tesouro = st.file_uploader("Upload da Planilha do Tesouro Direto", type=["xlsx", "csv"], key="up_tesouro")
        if arq_tesouro:
            try:
                if arq_tesouro.name.endswith('.csv'):
                    txt_tes = arq_tesouro.getvalue().decode('utf-8-sig', errors='ignore')
                    sep_tes = ',' if txt_tes.count(',') > txt_tes.count(';') else ';'
                    df_up_tes = pd.read_csv(io.StringIO(txt_tes), sep=sep_tes)
                else:
                    df_up_tes = pd.read_excel(arq_tesouro)
                
                mapeamento_tes = {}
                for col in df_up_tes.columns:
                    c_up = str(col).strip().upper()
                    if 'TÍTULO' in c_up or 'TITULO' in c_up: mapeamento_tes[col] = 'Título'
                    elif 'OPERAÇÃO' in c_up or 'COMPRA' in c_up: mapeamento_tes[col] = 'Data Compra'
                    elif 'BRUTO' in c_up or 'INVEST' in c_up: mapeamento_tes[col] = 'Valor Investido (R$)'
                    elif 'VENCIMENTO' in c_up or 'ANO' in c_up: mapeamento_tes[col] = 'Ano Vencimento'
                    elif 'RENTABILIDADE' in c_up or 'TAXA' in c_up: mapeamento_tes[col] = 'Rentabilidade_Raw'
                
                df_up_tes = df_up_tes.rename(columns=mapeamento_tes)
                
                if 'Rentabilidade_Raw' in df_up_tes.columns:
                    tipos, taxas = [], []
                    for val in df_up_tes['Rentabilidade_Raw'].astype(str):
                        val_up = val.upper()
                        if "SELIC" in val_up: tipos.append("Tesouro Selic")
                        elif "IPCA" in val_up: tipos.append("Pós-fixado (IPCA+)")
                        else: tipos.append("Pré-fixado")
                        
                        val_num = val.replace(',', '.')
                        nums = re.findall(r'\d+\.\d+|\d+', val_num)
                        taxas.append(float(nums[-1]) if nums else 0.0)
                        
                    df_up_tes['Tipo Taxa'] = tipos
                    df_up_tes['Taxa Contratada (%)'] = taxas

                for c in colunas_padrao:
                    if c not in df_up_tes.columns: df_up_tes[c] = None
                
                df_up_tes['Data Compra'] = pd.to_datetime(df_up_tes['Data Compra'], errors='coerce').dt.date
                df_up_tes['Valor Investido (R$)'] = df_up_tes['Valor Investido (R$)'].apply(limpar_numero)
                df_up_tes['Ano Vencimento'] = df_up_tes['Ano Vencimento'].apply(limpar_numero).astype(int)
                
                st.session_state.df_tesouro = df_up_tes[colunas_padrao]
                st.success("Planilha do Tesouro importada com sucesso!")
            except Exception as e:
                st.error(f"Erro ao processar planilha: {e}")

    df_calc = st.session_state.df_tesouro.copy()
    for c in colunas_padrao:
        if c not in df_calc.columns: df_calc[c] = None

    valores_atuais = []
    valores_futuros = []
    
    hoje = pd.Timestamp.now().date()
    
    ipca_proj = ipca_12m_hoje if ('ipca_12m_hoje' in locals() and ipca_12m_hoje > 0) else 4.0
    selic_atual = selic_hoje if ('selic_hoje' in locals() and selic_hoje > 0) else 10.5

    for _, row in df_calc.iterrows():
        try:
            inv = float(limpar_numero(row.get('Valor Investido (R$)', 0)))
            tx = float(limpar_numero(row.get('Taxa Contratada (%)', 0)))
            venc = int(limpar_numero(row.get('Ano Vencimento', hoje.year + 1)))
            
            dt_c = pd.to_datetime(row.get('Data Compra', hoje)).date()
            dias_decorridos = max(0, (hoje - dt_c).days)
            anos_decorridos = dias_decorridos / 365.25
            anos_vencimento = max(0.1, venc - hoje.year)
            
            tipo_t = str(row.get('Tipo Taxa')).strip()
            
            if tipo_t == "Pós-fixado (IPCA+)":
                tx_aplicada = tx + ipca_proj
            elif tipo_t == "Tesouro Selic":
                tx_aplicada = tx + selic_atual
            else:
                tx_aplicada = tx
                
            v_atual = inv * ((1 + (tx_aplicada/100)) ** anos_decorridos)
            v_futuro = v_atual * ((1 + (tx_aplicada/100)) ** anos_vencimento)
            
            valores_atuais.append(v_atual)
            valores_futuros.append(v_futuro)
        except:
            valores_atuais.append(0.0)
            valores_futuros.append(0.0)

    df_calc['Valor Atualizado (R$)'] = valores_atuais
    df_calc['Valor Futuro no Vencimento'] = valores_futuros
    
    df_calc['Valor Investido (R$)'] = df_calc['Valor Investido (R$)'].apply(limpar_numero)
    df_calc['Taxa Contratada (%)'] = df_calc['Taxa Contratada (%)'].apply(limpar_numero)
    df_calc['Ano Vencimento'] = df_calc['Ano Vencimento'].apply(limpar_numero).astype(int)

    st.markdown("#### 📝 Lançamentos e Projeções (Tabela Interativa)")
    
    df_editado_t = st.data_editor(
        df_calc,
        key="editor_tesouro_fix", 
        num_rows="dynamic",
        use_container_width=True,
        hide_index=True,
        column_order=["Título", "Data Compra", "Valor Investido (R$)", "Tipo Taxa", "Taxa Contratada (%)", "Ano Vencimento", "Valor Atualizado (R$)", "Valor Futuro no Vencimento"],
        column_config={
            "Tipo Taxa": st.column_config.SelectboxColumn("Tipo Taxa", options=["Pré-fixado", "Pós-fixado (IPCA+)", "Tesouro Selic"], required=True),
            "Data Compra": st.column_config.DateColumn("Data Compra", format="DD/MM/YYYY"),
            "Ano Vencimento": st.column_config.NumberColumn("Ano Venc.", format="%d"),
            "Valor Investido (R$)": st.column_config.NumberColumn("Valor Investido (R$)", format="%.2f"),
            "Taxa Contratada (%)": st.column_config.NumberColumn("Taxa / Prêmio (%)", format="%.2f"),
            "Valor Atualizado (R$)": st.column_config.NumberColumn("Valor Atualizado (Hoje)", disabled=True, format="R$ %.2f"),
            "Valor Futuro no Vencimento": st.column_config.NumberColumn("Projeção no Vencimento", disabled=True, format="R$ %.2f")
        }
    )
    
    if st.button("🔄 Salvar Edições Manuais da Tabela", type="secondary"):
        st.session_state.df_tesouro = df_editado_t[colunas_padrao]
        if st.session_state.username:
            salvar_dados_completos_db(st.session_state.username)
        # RERUN REMOVIDO DAQUI TAMBÉM.
        st.success("Alterações salvas permanentemente no banco de dados!")

    if not df_calc.empty:
        tot_investido = df_calc['Valor Investido (R$)'].sum()
        tot_atual = df_calc['Valor Atualizado (R$)'].sum()
        tot_projetado = df_calc['Valor Futuro no Vencimento'].sum()
        
        lucro_atual = tot_atual - tot_investido
        
        st.markdown("---")
        st.markdown("#### 📊 Resumo do Tesouro Direto")
        c_tot1, c_tot2, c_tot3 = st.columns(3)
        c_tot1.metric("Total Investido", f_brl(tot_investido))
        
        rent_ate_hoje = f_pct((tot_atual / tot_investido - 1) * 100) if tot_investido > 0 else "0.00%"
        c_tot2.metric("Valor Atualizado (Hoje)", f_brl(tot_atual), rent_ate_hoje)
        
        margem_futura = f_pct((tot_projetado / tot_atual - 1) * 100) if tot_atual > 0 else "0.00%"
        c_tot3.metric("Projeção no Vencimento", f_brl(tot_projetado), margem_futura)
        
with tab_extrato:
    st.markdown("### 📜 Histórico de Transações e Ganho de Capital")
    st.info("Registro cronológico imutável de todas as compras e vendas efetuadas no terminal.")
    
    if 'df_transacoes' in st.session_state and not st.session_state.df_transacoes.empty:
        df_tx_plot = st.session_state.df_transacoes.sort_values("Data", ascending=False)
        
        lucro_total = df_tx_plot[df_tx_plot['Tipo'] == 'Venda']['Resultado Realizado'].sum()
        
        c_card1, c_card2 = st.columns(2)
        if lucro_total >= 0:
            c_card1.metric("Ganho de Capital Líquido (Lucro Realizado)", f_brl(lucro_total))
        else:
            c_card1.metric("Prejuízo Consolidado Realizado", f_brl(lucro_total))
            
        st.markdown("#### 📑 Histórico de Ordens Lançadas")
        
        format_tx = {
            "Quantidade": lambda x: f"{int(x)}",
            "Preço Unitário": f_brl,
            "Preço Médio na Época": lambda x: f_brl(x) if x > 0 else "-",
            "Resultado Realizado": lambda x: f_brl(x) if x != 0 else "-"
        }
        st.dataframe(df_tx_plot.style.format(format_tx), use_container_width=True, hide_index=True)
        
        csv_tx = df_tx_plot.to_csv(index=False, sep=';', encoding='utf-8-sig')
        st.download_button(
            label="📥 Baixar Livro de Transações Completo (CSV Backup)",
            data=csv_tx,
            file_name=f"Extrato_Transacoes_{st.session_state.username}.csv",
            use_container_width=True
        )
    else:
        st.info("Nenhuma movimentação registrada no livro de transações.")

with tab_ia:
    st.markdown("### 💬 Comitê de IA Sênior")
    cb1, cb2 = st.columns(2)
    
    if cb1.button("🗑️ Limpar Chat e Iniciar Nova Sessão", use_container_width=True):
        st.session_state.historico_chat = [{"role": "assistant", "content": f"Saudações, {st.session_state.username}. O terminal foi limpo."}]
        st.rerun()
        
    if HAS_DOCX and len(st.session_state.historico_chat) > 1:
        doc_buffer = export_docx(st.session_state.historico_chat)
        cb2.download_button("📄 Exportar Conversa de IA em Documento (Word)", data=doc_buffer, file_name=f"Parecer_IA_{st.session_state.username}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
    elif not HAS_DOCX: 
        cb2.caption("⚠️ Instale a biblioteca 'python-docx' no seu arquivo requirements.txt para habilitar o botão de exportação em Word.")

    api_key_secreta = st.secrets.get("GEMINI_API_KEY", "")
    if not api_key_secreta: 
        api_key_secreta = st.text_input("Insira sua Gemini API Key (Apenas uma vez):", type="password")
        
    for m in st.session_state.historico_chat:
        with st.chat_message(m["role"]): 
            st.write(m["content"])
        
    if prompt := st.chat_input("Ex: 'A BBAS3 está sendo negociada com desconto?' ou 'Qual sua avaliação da minha carteira?'"):
        with st.chat_message("user"): 
            st.write(prompt)
            
        st.session_state.historico_chat.append({"role": "user", "content": prompt})
        
        with st.chat_message("assistant"):
            with st.spinner("Analisando cruzamentos de dados e o cenário macroeconômico..."):
                
                tot_b3 = df_perf_final['Saldo Atual'].sum() if ('df_perf_final' in locals() and not df_perf_final.empty) else 0.0
                try:
                    tot_tes = st.session_state.df_tesouro['Valor Investido (R$)'].apply(limpar_numero).sum() if (isinstance(st.session_state.df_tesouro, pd.DataFrame) and not st.session_state.df_tesouro.empty) else 0.0
                except:
                    tot_tes = 0.0
                    
                tot_geral = tot_b3 + tot_tes
                
                ctx_resumo = f"PATRIMÔNIO EXATO (NÃO CALCULE, USE ESSES VALORES CEGAMENTE): Total Consolidado = R$ {tot_geral:,.2f} | B3 (Ações/FIIs) = R$ {tot_b3:,.2f} | Tesouro Direto = R$ {tot_tes:,.2f}."
                
                ctx_c = str(st.session_state.dados_mercado) if st.session_state.dados_mercado else "O usuário não conectou a carteira B3 no momento."
                ctx_t = st.session_state.df_tesouro.to_dict(orient='records') if isinstance(st.session_state.df_tesouro, pd.DataFrame) and not st.session_state.df_tesouro.empty else "Sem títulos no Tesouro."
                ctx_m = f"Selic Vigente: {f_pct(selic_hoje)}|IPCA Atual 12m: {f_pct(ipca_12m_hoje)}. Projeções Focus para {ano_atual}: Selic {f_pct(proj_focus.get(f'Selic_{ano_atual}'))}/IPCA {f_pct(proj_focus.get(f'IPCA_{ano_atual}'))}"
                
                h_txt = "\n".join([f"{'Usuário' if h['role']=='user' else 'Gestora IA'}: {h['content']}" for h in st.session_state.historico_chat[-6:-1]])
                
                sys_p = (
                    f"Você é um Analista CNPI Sênior de alta performance. \n"
                    f"[{ctx_resumo}]\n"
                    f"[Carteira B3]: {ctx_c}\n"
                    f"[Carteira Tesouro Direto]: {ctx_t}\n"
                    f"[Cenário Macro]: {ctx_m}\n\n"
                    f"DIRETRIZ DE CONTINUIDADE: Use o HISTÓRICO RECENTE abaixo para não perder o fio da conversa com o cliente.\n"
                    f"REGRA ESTRITA: 1) NUNCA tente somar os valores JSON das carteiras por conta própria, pois LLMs erram na matemática de arrays. Use os totais passados em PATRIMÔNIO EXATO sempre que questionado sobre saldos totais. 2) Foque a sua análise na qualidade dos ativos, setores e estratégia de investimentos baseando-se no cenário macroeconômico atual.\n"
                    f"=== HISTÓRICO DA CONVERSA ===\n{h_txt}"
                )
                
                resposta_ia = "⚠️ Chave API da IA ausente ou incorreta."
                if api_key_secreta:
                    try:
                        import google.generativeai as genai
                        genai.configure(api_key=api_key_secreta)
                        sucesso_ia = False
                        erro_log = ""
                        
                        for mdl in ['gemini-2.5-flash', 'gemini-1.5-flash']:
                            try:
                                resposta_ia = genai.GenerativeModel(mdl).generate_content([sys_p, prompt]).text
                                sucesso_ia = True
                                break 
                            except Exception as e_ia: 
                                erro_log = str(e_ia)
                                continue
                                
                        if not sucesso_ia: 
                            resposta_ia = f"⚠️ Falha de comunicação de rede com os servidores da IA: {erro_log}"
                    except Exception as e_motor: 
                        resposta_ia = f"⚠️ Falta de dependência no servidor (motor genai): {e_motor}"
                        
                st.write(resposta_ia)
                
        st.session_state.historico_chat.append({"role": "assistant", "content": resposta_ia})
        
        if st.session_state.username:
            salvar_dados_completos_db(st.session_state.username) 
            
        st.rerun()
